from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path(r"D:\笛语跨品牌服装搭配专家内核\DIYU-CBFSK-PRD-v1.1-M0-执行申请_20260812.docx")

FONT_BODY = "Microsoft YaHei"
FONT_MONO = "Consolas"
INK = "172033"
BLUE = "22577A"
BLUE_DARK = "15384F"
MUTED = "667085"
LIGHT = "EEF4F8"
LIGHTER = "F7F9FB"
BORDER = "C7D3DD"
GREEN = "2F6B4F"
AMBER = "8A5A00"
RED = "9B1C1C"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_borders(table, color=BORDER, size=6) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), str(size))
        tag.set(qn("w:color"), color)


def set_table_geometry(table, widths_dxa: list[int], indent_dxa=120) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths_dxa[idx])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_font(run, name=FONT_BODY, size=10.5, color=INK, bold=None, italic=None) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def format_cell_text(cell, bold=False, color=INK, size=9.5, align=WD_ALIGN_PARAGRAPH.LEFT) -> None:
    for p in cell.paragraphs:
        p.alignment = align
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.08
        for r in p.runs:
            set_font(r, size=size, color=color, bold=bold)


def add_body(doc, text: str, *, bold_prefix: str | None = None, after=6, keep=False):
    p = doc.add_paragraph(style="Normal")
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.keep_together = keep
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_font(r1, bold=True)
        r2 = p.add_run(text[len(bold_prefix):])
        set_font(r2)
    else:
        r = p.add_run(text)
        set_font(r)
    return p


def add_bullet(doc, text: str, level=0, *, bold_prefix: str | None = None):
    style = "List Bullet" if level == 0 else "List Bullet 2"
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(4)
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_font(r1, bold=True)
        r2 = p.add_run(text[len(bold_prefix):])
        set_font(r2)
    else:
        r = p.add_run(text)
        set_font(r)
    return p


def add_number(doc, text: str):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    set_font(r)
    return p


def add_callout(doc, label: str, text: str, fill=LIGHT, color=BLUE_DARK):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360], indent_dxa=120)
    set_table_borders(table, color=BORDER, size=5)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(f"{label}\n")
    set_font(r, size=10, color=color, bold=True)
    r2 = p.add_run(text)
    set_font(r2, size=10.5, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_table(doc, headers: list[str], rows: list[list[str]], widths: list[int], aligns=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    set_table_borders(table)
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for idx, text in enumerate(headers):
        hdr.cells[idx].text = text
        set_cell_shading(hdr.cells[idx], LIGHT)
        format_cell_text(hdr.cells[idx], bold=True, color=BLUE_DARK, size=9.5,
                         align=(aligns[idx] if aligns else WD_ALIGN_PARAGRAPH.LEFT))
    for row_data in rows:
        row = table.add_row()
        for idx, text in enumerate(row_data):
            row.cells[idx].text = text
            if len(table.rows) % 2 == 1:
                set_cell_shading(row.cells[idx], LIGHTER)
            format_cell_text(row.cells[idx], size=9.2,
                             align=(aligns[idx] if aligns else WD_ALIGN_PARAGRAPH.LEFT))
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_code_block(doc, lines: list[str]):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    set_table_borders(table, color="D8DEE5", size=4)
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F5F7F9")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.05
    for idx, line in enumerate(lines):
        run = p.add_run(line + ("\n" if idx < len(lines) - 1 else ""))
        set_font(run, name=FONT_MONO, size=8.8, color="25313C")
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def set_cell_no_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_font(run, size=8.5, color=MUTED)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_end])


def configure_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT_BODY
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_BODY)
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.12

    settings = {
        "Title": (24, BLUE_DARK, 0, 5),
        "Subtitle": (12.5, MUTED, 0, 16),
        "Heading 1": (16, BLUE, 15, 7),
        "Heading 2": (13, BLUE, 11, 5),
        "Heading 3": (11.5, BLUE_DARK, 8, 4),
    }
    for name, (size, color, before, after) in settings.items():
        style = styles[name]
        style.font.name = FONT_BODY
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_BODY)
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = name not in ("Subtitle",)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Bullet 2", "List Number"):
        style = styles[name]
        style.font.name = FONT_BODY
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_BODY)
        style.font.size = Pt(10.5)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.12

    if "Small Meta" not in styles:
        meta = styles.add_style("Small Meta", WD_STYLE_TYPE.PARAGRAPH)
    else:
        meta = styles["Small Meta"]
    meta.font.name = FONT_BODY
    meta._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_BODY)
    meta.font.size = Pt(9)
    meta.font.color.rgb = RGBColor.from_string(MUTED)
    meta.paragraph_format.space_after = Pt(2)


def configure_sections(doc: Document) -> None:
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(0.82)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        section.header_distance = Inches(0.35)
        section.footer_distance = Inches(0.38)


def build_document() -> Document:
    doc = Document()
    configure_styles(doc)
    configure_sections(doc)

    section = doc.sections[0]
    header_p = section.header.paragraphs[0]
    header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hr = header_p.add_run("DIYU-CBFSK-001  |  EXECUTION APPLICATION")
    set_font(hr, size=8.5, color=MUTED, bold=True)
    add_page_number(section.footer.paragraphs[0])

    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_before = Pt(18)
    kicker.paragraph_format.space_after = Pt(8)
    kr = kicker.add_run("CLAUDE CODE APPROVAL REQUEST")
    set_font(kr, size=9.5, color=BLUE, bold=True)

    title = doc.add_paragraph(style="Title")
    tr = title.add_run("笛语跨品牌服装搭配专家内核")
    set_font(tr, size=24, color=BLUE_DARK, bold=True)
    subtitle = doc.add_paragraph(style="Subtitle")
    sr = subtitle.add_run("PRD v1.1 修复与 M0 合同冻结执行申请")
    set_font(sr, size=12.5, color=MUTED)

    add_table(
        doc,
        ["控制项", "内容"],
        [
            ["申请 ID", "DIYU-CBFSK-WORK-EXEC-REQUEST-20260812-001"],
            ["交接基线", "DIYU-CBFSK-WORK-HANDOFF-20260812-001"],
            ["申请日期", "2026-08-12"],
            ["审批对象", "Claude Code - Contract / Completeness Guardian"],
            ["业务与产品最终裁决", "Founder"],
            ["当前状态", "PROJECT_INITIATED / EXECUTION_NOT_STARTED"],
            ["申请性质", "文档修复与 M0 合同基础；不是知识蒸馏、生产接入或 Serving 发布"],
        ],
        [1900, 7460],
    )

    add_callout(
        doc,
        "申请结论",
        "请 Claude Code 对本申请的范围完整性、交付物闭环、越界风险、清单一致性与可验收性进行独立审查，输出 APPROVE / APPROVE_WITH_CONDITIONS / REJECT。Claude Code 的结论是工程与完整性审查意见，不替代 Founder 的最终业务授权。",
    )

    doc.add_heading("一、申请执行的唯一范围", level=1)
    add_body(doc, "本次申请只允许两步连续交付：")
    add_number(doc, "先形成 PRD v1.1 Delta，将新增 Founder 裁决、完整性审查的有效修复项和独立解耦边界写回正式产品合同。")
    add_number(doc, "PRD v1.1 通过 Founder 裁决后，只执行 GKB-CROSS-BRAND-STYLING-EXPERT-KERNEL-M0-CONTRACT-AND-BENCHMARK-FOUNDATION-001，完成 M0 合同冻结并为 M1-M2 建立可执行基础。")
    add_body(doc, "本申请不授权任何其他里程碑施工。")

    doc.add_heading("二、Founder 已冻结裁决", level=1)
    decisions = [
        ["项目关系", "STANDALONE_DECOUPLED", "独立代码/项目空间、合同、数据模型、评测、版本、部署、运维与商业路线；不依赖笛语现有底座状态。"],
        ["未来集成", "Diyu foundation consumes kernel", "笛语底座只能通过版本化 API / SDK / Schema / 组件包消费本项目能力，不得拆散内核或重新定义其知识与决策合同。"],
        ["默认建设预算", "125-185 人月 / 15-24 个月", "作为独立完整建设的默认基线；仅在未来验证中立基础设施可安全复用时重算增量。"],
        ["专家内核", "One Universal Expert Kernel", "品类通过 Category Adapter 提供特有语义、功能和安全约束；品牌只提供 Brand Truth Pack，不复制品牌专属专家。"],
        ["生产模型策略", "Provider-agnostic, API-swappable", "系统必须能绑定和切换任意符合合同的模型；测试期按需提出 API 需求，由 Founder 提供凭证；本阶段不考虑自部署。"],
        ["商品属性真源", "Brand-specific database", "品牌专属数据库是商品属性、款色码与品牌事实的授权真源；模型推断只能作为候选补充，不得覆盖真源。"],
        ["发布终态", "Human-in-the-loop only", "本项目不考虑自动发布终态；所有内容均需有权人工审阅与显式发布动作。"],
        ["商业 KPI", "Founder judgment", "系统记录事先冻结的业务信号和数据，但不设代替 Founder 的自动商业通过分数；是否继续、扩大或停止由 Founder 综合裁决。"],
    ]
    add_table(doc, ["裁决项", "冻结值", "执行含义"], decisions, [1700, 2500, 5160])

    doc.add_heading("三、个体资料、照片与顾客档案合同", level=1)
    add_body(doc, "系统同时支持“商品适配判断”与“具体顾客档案”，两者不得混写。")
    add_table(
        doc,
        ["模式", "输入与保存", "约束"],
        [
            ["商品适配知识", "系统对每一款商品适合的体型、肤色、气质、发型、预算、舒适边界、场景等形成可审查的专业判断。", "属于商品/专家知识，不是某个真实顾客的个人档案；不得用未接受候选知识正式服务。"],
            ["非实名/一次性适配", "用户可自然语言自述、简短问答、导购代录或上传照片。默认不长期保存姓名和原始肖像；可保存为本次搭配必要的结构化特征。", "遵守最小必要、目的限定、保存期限与可删除。去标识化后如仍可关联个人，仍按个人信息管理。"],
            ["完整顾客档案", "当顾客、品牌或门店业务明确需要长期顾客档案时，系统应支持完整保存姓名、联系信息、原始照片、身体与审美特征、预算、舒适边界、尺码、试穿/购买/搭配与反馈记录。", "必须经授权角色发起并取得顾客有效授权；应支持查询、更正、撤回、删除与权限审计。未满14周岁顾客需监护人同意和专门规则。"],
        ],
        [1800, 4200, 3360],
    )
    add_callout(
        doc,
        "合规边界",
        "Founder 已裁决产品必须支持完整顾客档案；M0 必须建立 compliance_review_contract 并指定责任人。具体同意文本、保存期限、照片处理、跨境传输与未成年人规则须经法务/合规核验，不由执行代理自行宣布满足法律要求。",
        fill="FFF7E8",
        color=AMBER,
    )

    doc.add_heading("四、Pilot 三品类与夹具品牌裁决", level=1)
    add_body(doc, "本轮裁决将“Pilot先选三个真实品类”改为“三类品类的合成夹具品牌验证”。三类不需在本申请中指定具体名称，应在 M2 依据覆盖矩阵冻结，且至少包含：")
    add_bullet(doc, "一个成人审美与身体适配类；")
    add_bullet(doc, "一个未成年人或家庭安全类；")
    add_bullet(doc, "一个功能优先类。")
    add_table(
        doc,
        ["项目", "冻结合同"],
        [
            ["夹具品牌供给", "Codex 依据已冻结 Schema 和品类合同模拟合成品牌事实包、模拟货盘与任务样例。"],
            ["审查", "Claude Code 在独立工作区审查 Schema 完整性、内部自洽性、边界覆盖、品类安全、可重放性与泄漏风险；不审批其为真实品牌事实。"],
            ["资产标识", "必须显式标记 synthetic_fixture_brand / non_real_inventory / non_commercial_evidence，不得使用可误认为真实品牌的名称或叙事。"],
            ["可证明范围", "可验证 Schema、Checker、失败关闭、决策轨迹、迁移方法与工程可运行性。"],
            ["不可证明范围", "不得作为真实品牌适配、真实库存、门店采用、销售结果、Commercial V1.0 或专家商业能力的证据。"],
            ["时序", "M0 只冻结夹具品牌方法和责任；M1 Schema 稳定后才允许合成；M2 冻结隔离的评测分割。"],
        ],
        [1900, 7460],
    )
    add_callout(
        doc,
        "关键事实边界",
        "“夹具品牌验证”是开发和离线评测，不是 M11 所定义的真实多品牌、多门店、真实库存商业试点。如未来不做真实品牌试点，则 M11 应更名为“合成夹具试验”，且不得对外宣称已通过商业试点门。",
        fill="FDECEC",
        color=RED,
    )

    doc.add_heading("五、PRD v1.1 必须修复的执行合同", level=1)
    repairs = [
        ["P0-01", "独立解耦和预算基线", "写入 standalone_decoupled、standalone_full_stack、未来消费方向与125-185人月/15-24个月；删除“默认复用底座90-135人月”冲突措辞。"],
        ["P0-02", "M0唯一清单", "13、14、17.1三处完全对齐，纳入knowledge_state_machine、non_goals_and_stop_conditions、architecture_and_integration_boundary、compliance_review_contract。"],
        ["P0-03", "M1 Schema收口", "纳入BrandTruthPack与全部正式输入/输出对象Schema，或一个绑定所有对象的可验证bundle。"],
        ["P0-04", "合规责任与发布门", "M0建立合规待核验清单与责任人；新增独立Safety, Privacy & Compliance Gate；明确永久人工发布。"],
        ["P0-05", "夹具品牌/语料工作流", "冻结Codex合成、Claude Code独立完整性审查、合成资产标记、隔离规则、可证明/不可证明边界和人月/调用成本记录。"],
        ["P0-06", "投资阶梯与关键路径", "不得继续使用24-36人月POC与M0-M12无映射的口径；M0输出关键路径、限额POC证据包及continue/pivot/stop裁决点。"],
        ["P1-01", "M2评测治理三件套", "rater_calibration、sampling_design、benchmark_revision_protocol；修订不得由失败触发降标，且需Founder裁决、版本增量与重跑。"],
        ["P1-02", "承诺-功能-组件闭环", "新增单品多套生命FR；明确门店成交辅助输出；将跨品牌/品类比较并入迁移验收；由M6组件显式承接StylingDiagnosis。"],
        ["P1-03", "反馈隔离与运营治理", "落地feedback_to_candidate_ledger；内容表现、门店修改与顾客反馈不得直写专家知识真源。"],
        ["P1-04", "模型和API运行合同", "提供商中立适配层、模型能力清单、回归、失败关闭、凭证不落盘/不输出、调用成本与分场景时延的延迟冻结节点。"],
        ["P1-05", "商业信号与Founder裁决", "M11前冻结业务信号名录和数据口径；不设自动KPI通过线；最终继续/扩大/停止由Founder综合裁决。"],
    ]
    add_table(doc, ["编号", "修复项", "最小验收合同"], repairs, [900, 2200, 6260])

    doc.add_heading("六、M0 唯一交付清单", level=1)
    add_body(doc, "PRD v1.1 通过 Founder 裁决后，M0 在第13节、第14节和第17.1节必须使用以下同一清单，不得使用包含关系不明的两套清单：")
    m0_files = [
        "project_charter.v1.1.yaml",
        "capability_contract.v1.1.yaml",
        "category_scope_contract.v1.1.yaml",
        "input_output_boundary.v1.1.yaml",
        "role_and_decision_rights.v1.1.yaml",
        "knowledge_state_machine.v1.1.yaml",
        "non_goals_and_stop_conditions.v1.1.yaml",
        "architecture_and_integration_boundary.v1.1.yaml",
        "model_provider_and_api_boundary.v1.1.yaml",
        "individual_data_and_customer_record_contract.v1.1.yaml",
        "product_attribute_source_contract.v1.1.yaml",
        "human_publish_only_contract.v1.1.yaml",
        "synthetic_fixture_brand_governance.v1.1.yaml",
        "compliance_review_contract.v1.1.yaml",
        "investment_stage_and_critical_path_brief.v1.1.md",
        "M1_object_model_and_schema_brief.v1.1.md",
        "M2_benchmark_freeze_brief.v1.1.md",
        "M0 checker / fixtures / report / receipt / replay manifest",
    ]
    for item in m0_files:
        add_bullet(doc, item)
    add_body(doc, "以上文件名允许根据项目目录做 delta 适配，但语义对象、负责人、失败状态与验收范围不得缺失。")

    doc.add_heading("七、本次明确禁止事项", level=1)
    prohibitions = [
        "不得开始大规模GPT x Claude Code知识蒸馏、批量候选生成或专家知识接受。",
        "不得在M1 Schema冻结前正式生成夹具品牌资产；M0只定义其治理合同。",
        "不得把夹具品牌、模拟库存或合成顾客当作真实业务证据。",
        "不得接入真实品牌生产、正式库存、真实顾客敏感数据或任何内容发布账号。",
        "不得创建或翻转Serving、RAG、DIFY、生产发布或Commercial V1.0状态。",
        "不得为任何品牌复制独立专家内核、补写隐藏品牌专属Prompt或向生成链泄漏隐藏答案。",
        "不得获取、输出或持久化Founder提供的模型API明文凭证；本次也不得以“测试方便”为由索要凭证。",
        "不得将Claude Code的工程审批写成Founder的业务批准，也不得自行更改本申请中已冻结的产品裁决。",
    ]
    for item in prohibitions:
        add_bullet(doc, item)

    doc.add_heading("八、执行与审查分工", level=1)
    roles = [
        ["Founder", "冻结业务范围、风险接受、PRD v1.1、M0和商业延续性；提供测试期模型API但不提供凭证管理授权。"],
        ["Codex Engineering", "按批准合同修订PRD v1.1、落盘M0对象、Schema brief、Checker、Fixtures、Report、Receipt和重放证据；M1后才合成夹具品牌。"],
        ["Claude Code Guardian", "使用隔离上下文审查范围一致性、对象完整性、清单闭环、事实边界、隐藏评测泄漏、夹具自洽性和验收可重放性；不产生业务真源。"],
        ["外部专家/法务", "在后续指定节点审查专业正确性、品类安全与个人信息合规；本次申请不宣称已取得该等审查。"],
    ]
    add_table(doc, ["角色", "责任与边界"], roles, [2100, 7260])

    doc.add_heading("九、通过标准与审批回执", level=1)
    add_body(doc, "Claude Code 应逐项返回审查结论，而非只回复“同意”。")
    approval_rows = [
        ["范围一致", "是否严格停留在PRD v1.1 + M0，且未提前生成知识、夹具或生产能力。", "PASS / FAIL + 证据"],
        ["清单闭环", "M0唯一清单、M1 Schema、M2评测治理、负责人和失败状态是否可追溯。", "PASS / FAIL + 缺口"],
        ["边界安全", "独立解耦、API凭证、顾客档案、人工发布和夹具品牌不可证明范围是否清楚。", "PASS / FAIL + 风险"],
        ["可验收性", "交付物是否可运行Checker、可重放、可出Report/Receipt，不以文档数量替代结果。", "PASS / FAIL + 修复建议"],
        ["最终意见", "上述问题是否已达到可向Founder申请执行授权的程度。", "APPROVE / APPROVE_WITH_CONDITIONS / REJECT"],
    ]
    add_table(doc, ["审查项", "要求", "回执格式"], approval_rows, [1600, 5460, 2300])
    add_code_block(doc, [
        "approval_request_id: DIYU-CBFSK-WORK-EXEC-REQUEST-20260812-001",
        "reviewer_role: claude_code_contract_completeness_guardian",
        "decision: APPROVE | APPROVE_WITH_CONDITIONS | REJECT",
        "scope_alignment: PASS | FAIL",
        "deliverable_closure: PASS | FAIL",
        "fact_and_safety_boundaries: PASS | FAIL",
        "blocking_findings: []",
        "non_blocking_findings: []",
        "required_delta_before_execution: []",
        "reviewed_version_or_file_hash: <required>",
        "review_timestamp: <required>",
    ])

    doc.add_heading("十、批准后的唯一后续动作", level=1)
    add_callout(
        doc,
        "唯一下一步",
        "若 Claude Code 返回 APPROVE 或无阻断性的 APPROVE_WITH_CONDITIONS，则将其回执与本申请一并报 Founder。只有 Founder 明确批准后，才允许落盘 PRD v1.1 Delta；PRD v1.1 再经 Founder 裁决后，才允许执行 M0。任何 REJECT 或阻断性条件均必须先修复并重审。",
        fill="EAF5EF",
        color=GREEN,
    )

    add_body(doc, "文档结束。", after=0)

    # Prevent table rows from splitting where practical.
    for table in doc.tables:
        for row in table.rows:
            set_cell_no_split(row)

    return doc


def main() -> None:
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = build_document()
    doc.core_properties.title = "笛语跨品牌服装搭配专家内核 PRD v1.1 与 M0 执行申请"
    doc.core_properties.subject = "Claude Code 合同与完整性审批"
    doc.core_properties.author = "Codex Engineering Workspace"
    doc.core_properties.keywords = "DIYU-CBFSK, PRD v1.1, M0, Claude Code, approval"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
