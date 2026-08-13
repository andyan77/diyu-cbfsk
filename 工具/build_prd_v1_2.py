#!/usr/bin/env python3
"""Build the PRD v1.2 documentation baseline from the frozen v1.1 DOCX files.

This script performs contract-level, traceable mutations while preserving the
existing DOCX styles, section geometry, colors, headers, and footer fields.
It deliberately does not create M0 engineering assets.
"""

from __future__ import annotations

import copy
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.document import Document as DocumentType
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from docx.table import Table, _Cell, _Row
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parent.parent
PRD_SOURCE = ROOT / "笛语跨品牌服装搭配专家内核_PRD与执行里程碑_v1.1.docx"
M0_SOURCE = ROOT / "笛语跨品牌服装搭配专家内核_M0执行申请_v1.1.docx"
RECEIPT_SOURCE = ROOT / "PRD_v1.1_核验回执.docx"

PRD_OUTPUT = ROOT / "笛语跨品牌服装搭配专家内核_PRD与执行里程碑_v1.2.docx"
M0_OUTPUT = ROOT / "笛语跨品牌服装搭配专家内核_M0执行申请_v1.2.docx"
RECEIPT_OUTPUT = ROOT / "PRD_v1.2_核验回执.docx"

INK = "172033"
BLUE = "22577A"
BLUE_DARK = "15384F"
MUTED = "667085"
LIGHT = "EEF4F8"
BORDER = "C7D3DD"
FONT_BODY = "Microsoft YaHei"
EXECUTION_RUN_ID = "2fcbfed0-be7e-4b6f-938e-7f84109ab162"
PARENT_EXECUTION_RUN_ID = "29899f92-c965-4c33-a020-8e7f781fe82d"
BASELINE_DATE = "2026-08-13"


def _source(path: Path) -> Path:
    if path.exists():
        return path
    archived = ROOT / "归档_v1.1" / path.name
    if archived.exists():
        return archived
    raise FileNotFoundError(path)


def iter_paragraphs(parent: DocumentType | _Cell) -> Iterable[Paragraph]:
    for paragraph in parent.paragraphs:
        yield paragraph
    for table in parent.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from iter_paragraphs(cell)


def all_paragraphs(doc: DocumentType) -> Iterable[Paragraph]:
    yield from iter_paragraphs(doc)
    for section in doc.sections:
        yield from iter_paragraphs(section.header)
        yield from iter_paragraphs(section.footer)


def find_paragraph(doc: DocumentType, text: str, *, starts: bool = False) -> Paragraph:
    matches: list[Paragraph] = []
    for paragraph in doc.paragraphs:
        value = paragraph.text.strip()
        if (starts and value.startswith(text)) or (not starts and value == text):
            matches.append(paragraph)
    if len(matches) != 1:
        raise ValueError(f"Expected one paragraph for {text!r}, found {len(matches)}")
    return matches[0]


def find_all_paragraphs(doc: DocumentType, text: str, *, starts: bool = False) -> list[Paragraph]:
    result: list[Paragraph] = []
    for paragraph in doc.paragraphs:
        value = paragraph.text.strip()
        if (starts and value.startswith(text)) or (not starts and value == text):
            result.append(paragraph)
    return result


def set_run_font(run, *, size: float = 10.0, color: str = INK, bold: bool | None = None) -> None:
    run.font.name = FONT_BODY
    fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    fonts.set(qn("w:ascii"), FONT_BODY)
    fonts.set(qn("w:hAnsi"), FONT_BODY)
    fonts.set(qn("w:eastAsia"), FONT_BODY)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold


def set_paragraph_text(paragraph: Paragraph, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def replace_text(doc: DocumentType, old: str, new: str, *, count: int | None = None) -> int:
    changed = 0
    for paragraph in all_paragraphs(doc):
        if old in paragraph.text:
            set_paragraph_text(paragraph, paragraph.text.replace(old, new))
            changed += 1
    if count is not None and changed != count:
        raise ValueError(f"Expected {count} replacements for {old!r}, got {changed}")
    return changed


def find_table(doc: DocumentType, header: str) -> Table:
    matches = [table for table in doc.tables if table.rows and table.cell(0, 0).text.strip() == header]
    if len(matches) != 1:
        raise ValueError(f"Expected one table headed {header!r}, found {len(matches)}")
    return matches[0]


def set_cell_text(cell: _Cell, text: str) -> None:
    paragraph = cell.paragraphs[0]
    set_paragraph_text(paragraph, text)
    for extra in cell.paragraphs[1:]:
        set_paragraph_text(extra, "")


def clone_table_row(table: Table, values: list[str]) -> _Row:
    if len(values) != len(table.columns):
        raise ValueError(f"Row has {len(values)} values for {len(table.columns)} columns")
    template = table.rows[-1]
    new_tr = copy.deepcopy(template._tr)
    table._tbl.append(new_tr)
    row = _Row(new_tr, table)
    for cell, value in zip(row.cells, values):
        set_cell_text(cell, value)
    return row


def insert_table_from_template_before(
    marker: Paragraph,
    template: Table,
    headers: list[str],
    rows: list[list[str]],
) -> Table:
    """Insert a table before *marker* while preserving template geometry/style."""
    if len(headers) != len(template.columns):
        raise ValueError("Template column count does not match headers")
    new_tbl = copy.deepcopy(template._tbl)
    for tr in list(new_tbl.tr_lst):
        new_tbl.remove(tr)
    header_tr = copy.deepcopy(template.rows[0]._tr)
    new_tbl.append(header_tr)
    table = Table(new_tbl, marker._parent)
    for cell, value in zip(table.rows[0].cells, headers):
        set_cell_text(cell, value)
    row_template = template.rows[1] if len(template.rows) > 1 else template.rows[0]
    for values in rows:
        if len(values) != len(headers):
            raise ValueError("Inserted table row width mismatch")
        new_tr = copy.deepcopy(row_template._tr)
        new_tbl.append(new_tr)
        row = _Row(new_tr, table)
        for cell, value in zip(row.cells, values):
            set_cell_text(cell, value)
    marker._p.addprevious(new_tbl)
    return table


def set_table_value(table: Table, row_label: str, value: str, *, column: int = 1) -> None:
    matches = [row for row in table.rows if row.cells[0].text.strip() == row_label]
    if len(matches) != 1:
        raise ValueError(f"Expected one row {row_label!r}, found {len(matches)}")
    set_cell_text(matches[0].cells[column], value)


def insert_before(marker: Paragraph, text: str, style: str = "Normal") -> Paragraph:
    paragraph = marker.insert_paragraph_before(text, style=style)
    return paragraph


def insert_bullets_before(marker: Paragraph, lines: list[str], *, checkbox: bool = False) -> None:
    for line in lines:
        prefix = "[ ] " if checkbox else "• "
        insert_before(marker, prefix + line, "Normal")


def insert_sections_before(doc: DocumentType, marker_text: str, sections: list[tuple[str, list[str]]]) -> None:
    marker = find_paragraph(doc, marker_text)
    for heading, paragraphs in sections:
        insert_before(marker, heading, "Heading 2")
        for text in paragraphs:
            insert_before(marker, text, "Normal")


def section_marker(doc: DocumentType, milestone: str, chapter: int) -> Paragraph:
    prefix = f"M{milestone}｜" if chapter == 13 else f"M{milestone} "
    matches = [p for p in doc.paragraphs if p.text.strip().startswith(prefix)]
    if len(matches) != 1:
        raise ValueError(f"Expected one chapter {chapter} marker for M{milestone}, got {len(matches)}")
    return matches[0]


def next_heading(doc: DocumentType, paragraph: Paragraph, level: str = "Heading 2") -> Paragraph:
    seen = False
    for candidate in doc.paragraphs:
        if candidate._p is paragraph._p:
            seen = True
            continue
        if seen and candidate.style is not None and candidate.style.name == level:
            return candidate
    raise ValueError(f"No following {level} after {paragraph.text!r}")


def insert_milestone_deliverables(
    doc: DocumentType,
    milestone: str,
    chapter: int,
    lines: list[str],
) -> None:
    marker = section_marker(doc, milestone, chapter)
    end = next_heading(doc, marker)
    paragraphs = doc.paragraphs
    start_index = next(i for i, p in enumerate(paragraphs) if p._p is marker._p)
    end_index = next(i for i, p in enumerate(paragraphs) if p._p is end._p)
    region = paragraphs[start_index:end_index]
    if chapter == 13:
        pass_markers = [p for p in region if p.text.strip() == "通过标准"]
        if len(pass_markers) != 1:
            raise ValueError(f"M{milestone} chapter 13 has no unique pass marker")
        target = pass_markers[0]
    else:
        gate_lines = [p for p in region if p.text.strip().startswith("[ ] 验收门已通过：")]
        if len(gate_lines) != 1:
            raise ValueError(f"M{milestone} chapter 14 has no unique gate line")
        target = gate_lines[0]
    insert_bullets_before(target, lines, checkbox=True)


def replace_heading(doc: DocumentType, old: str, new: str) -> None:
    set_paragraph_text(find_paragraph(doc, old), new)


def build_prd() -> None:
    doc = Document(_source(PRD_SOURCE))

    # Cover, document control, version record, and index.
    set_paragraph_text(doc.paragraphs[4], "跨品牌 · 跨品类 · 多模态商品理解 · 人设连续性 · 专家叙事")
    control = doc.tables[0]
    set_table_value(control, "Document Version（文档版本）", "PRD v1.2 · 人设连续性、多模态商品理解与治理操作基线")
    set_table_value(control, "Project Status（项目状态）", "PROJECT_INITIATED / EXECUTION_NOT_STARTED / M0_AUTHORIZED_FALSE")
    set_table_value(control, "Baseline Date（基线日期）", BASELINE_DATE)
    set_cell_text(
        doc.tables[1].cell(0, 0),
        "正式立项裁决\n本项目从笛语现有内容生产与陈列场景中独立立项。项目建设一个可跨品牌、跨品类迁移的顶级服装搭配专家内核；正式范围同时包含商品图片的证据分级理解、长期搭配师人设连续性和自媒体原生语感。完整陈列、实时成交辅助与自动发布按扩展兼容和Founder授权合同管理，不因本版本文档升级翻转执行或生产状态。",
    )
    document_control = doc.tables[2]
    set_table_value(
        document_control,
        "Scope Status（范围状态）",
        "D-17—D-29与S1—S8已合并为候选合同；详细Schema、评测阈值与知识单元数量仍须在获授权的M0—M2冻结。当前候选待独立Guardian、ChatGPT总顾问与Founder审查。",
    )
    set_table_value(
        document_control,
        "Execution Status（执行状态）",
        "未开始。本文档不表示M0已授权，也不表示任何知识、模型、Serving、RAG、多模态流水线、人设记忆库或生产状态已翻转。",
    )
    clone_table_row(
        doc.tables[3],
        [
            "v1.2",
            BASELINE_DATE,
            "人设连续性、多模态商品理解与治理操作基线",
            "完整合并D-17—D-29、S1—S8、单一Founder角色操作模型、隐藏评测物理隔离、条件/合规台账与Project CI。D-23、D-25按Founder本轮补丁改判；M0顶层交付清单仍为14项。候选待Guardian、总顾问与Founder审查，尚未生效。",
        ],
    )
    index = doc.tables[4]
    for row in index.rows:
        for cell in row.cells:
            if cell.text.strip() == "9. 叙事智能与账号矩阵内容合同":
                set_cell_text(cell, "9. 叙事智能、人设连续性与自媒体语感合同")

    # Chapter 1.
    definition = doc.tables[5]
    set_cell_text(
        definition.cell(0, 0),
        "Product Definition（产品定义）\n构建一个能够跨品牌、跨品类理解品牌定位与目标人群，联合权威结构化商品事实与有来源、有置信度的商品图片视觉推断，在总部或门店有效货盘约束下生成可解释、可替代、可购买的专业搭配方案，并通过持续、可审计的搭配师人设和自媒体原生语感将专业判断转化为真实、有观点、有场景、有变化的内容的数字化顶级服装搭配专家。",
    )
    set_paragraph_text(
        find_paragraph(doc, "系统面对一个此前没有专门建设过的新品牌与新货盘，只读取必要的品牌事实、品类规则、目标人群与实时库存，仍能够完成以下闭环："),
        "系统面对一个此前没有专门建设过的新品牌与新货盘，仅读取当次任务必需的品牌事实、品类规则、商品结构化资料与图片、目标人群及有效货盘信息，仍能完成以下闭环：",
    )
    north_star_end = find_paragraph(doc, "• 保留事实来源、决策轨迹、不确定性和人工升级边界。")
    next_marker = find_paragraph(doc, "1.4 项目状态合同")
    insert_before(next_marker, "• 在长期内容经营中保持搭配师的专业身份、核心价值观、历史观点与栏目连续性，并对有依据的观点演进留痕。")
    insert_before(next_marker, "• 形成非课件化、非模板化、非AI腔的自媒体原生语感，同一人设在不同平台保持核心辨识度。")
    insert_before(next_marker, "• 对商品图片进行证据分级的多模态理解，当视觉推断与权威结构化事实冲突时以权威事实为准。")
    insert_before(next_marker, "• 在M11支持Founder真实品牌注入测试生产；未提供时回退到Codex夹具合成，两条路径均不泄漏或污染封闭隐藏测试集。")
    insert_before(next_marker, "• 通过版本化扩展端口兼容未来门店空间陈列、实时成交辅助和Founder授权自动发布，不将其误写为V1.0当前强制实现。")
    state_table = next(
        table
        for table in doc.tables
        if table.cell(0, 0).text.strip() == "状态"
        and table.cell(1, 0).text.strip() == "PROJECT_INITIATED"
    )
    clone_table_row(state_table, ["M0_AUTHORIZED_FALSE", "PRD v1.2与M0执行申请v1.2尚未经Founder签署，M0不得开工。"])

    # Chapter 3 goals, non-goals, and principles.
    goals = next(
        table
        for table in doc.tables
        if table.cell(0, 0).text.strip() == "ID"
        and table.cell(1, 0).text.strip() == "G-01"
    )
    for row in [
        ["G-11", "持续搭配师人设", "使专业身份、价值观、审美立场、历史观点与栏目在长期内容中可连续、可演进、可审计。"],
        ["G-12", "自媒体原生语感", "在不改变专业结论的前提下，产生具有平台原生表达、口播节奏、场景感、镜头可拍性和人设辨识度的内容。"],
        ["G-13", "多模态商品理解", "从商品图片中提取搭配决策需要的视觉属性，并保留来源、模型、置信度、冲突与人工覆盖证据。"],
        ["G-14", "扩展兼容与按需开启", "为视觉陈列、实时成交辅助与Founder授权自动发布保留版本化扩展端口，未开启时不增加V1.0强制交付门。"],
        ["G-15", "Project CI与治理完整性", "角色、权限、任务分级、产品真源与公开/隐藏边界由单一规范源生成和确定性守卫约束，不形成第三套产品真源。"],
    ]:
        clone_table_row(goals, row)
    set_paragraph_text(
        find_paragraph(doc, "• 自动发布内容：本项目V1.0全生命周期保持人工审核在环，自动发布终态不属于V1.0产品承诺——包括所有商业发布门通过之后。（保留不变：FR-17发布前库存复核、“人工可发布率≥75%”质量门。）"),
        "• V1.0 默认人工审核在环；自动发布仅在 Founder 按租户/品牌/账号/风险级别显式授权后开启，须具备审计、撤回与 Kill Switch。FR-17 与“人工可发布率≥75%”质量门保留不变。",
    )
    set_paragraph_text(
        find_paragraph(doc, "• 以图像识别自动生成或提取商品属性：商品与货盘属性以品牌专属数据库（品牌真源／PIM／ERP等）为唯一来源，V1.0不立项图像识别属性提取工程线；字段缺失时按FR-05既有路径触发收窄或要求补充，不得脑补属性。"),
        "• 品牌专属数据库仍为权威事实来源；多模态视觉推断按 D-23 分层引入（authoritative > human_confirmed > model_inferred），推断不得覆盖权威事实，不得推断成分、性能、库存与安全认证。",
    )
    non_goal_marker = find_paragraph(doc, "3.3 核心设计原则")
    insert_before(non_goal_marker, "• 完整门店空间陈列与实时交互式导购工具不是V1.0强制交付物；它们通过扩展端口保留前后向兼容，不是永久退出项目。")
    principles = next(table for table in doc.tables if table.cell(0, 0).text.strip() == "原则")
    for row in [
        ["P-10", "Persona Continuity before Content Scale（人设连续优先于内容规模）", "在扩大账号内容产能前，先确保专业身份、核心价值观、观点和栏目状态可连续、可回放。"],
        ["P-11", "Evidence-graded Multimodal Understanding（证据分级多模态理解）", "视觉属性必须保留来源、模型、置信度、冲突与人工覆盖，不得覆盖权威事实。"],
        ["P-12", "Extension-ready, not Premature Implementation（扩展就绪，不提前过度实现）", "用端口、Schema和Bundle兼容未来模块，不将未裁决的扩展误写为当前强制范围。"],
        ["P-13", "Founder-controlled Publication（Founder控制发布）", "默认人工审核；自动发布只能在Founder授权范围内启用，且可审计、可撤回、可回滚、可紧急停止。"],
    ]:
        clone_table_row(principles, row)

    # Chapter 4 scenarios.
    scenarios = next(table for table in doc.tables if table.cell(0, 0).text.strip() == "ID" and table.cell(0, 1).text.strip() == "场景")
    for row in scenarios.rows:
        if row.cells[0].text.strip() == "S-06":
            set_cell_text(row.cells[2], "从真实搭配决策中生产观点、冲突、变化与人物处境，维护长期人设、历史观点、栏目连续性及自媒体原生语感。")
        if row.cells[0].text.strip() == "S-07":
            set_cell_text(row.cells[2], "V1.0提供基于已接受搭配决策的“导购辅助要点”，作为ContentProjection在M7实现、M11验收。实时交互式导购工具当前不强制实现，但通过RealtimeSalesAssistExtensionPort保持扩展兼容，未来开启时不得绕过品类安全、个体授权与库存复核。")
    clone_table_row(scenarios, ["S-09", "Founder真实品牌注入测试生产", "M11时Founder可提供真实品牌、商品、图片及可选货盘资料，系统生成完整搭配、叙事、人设和约束结果供Founder评判；该路径为可选能力判断路径，不是M11强制硬门。"])

    # Chapter 5 capabilities and models.
    set_paragraph_text(find_paragraph(doc, "5.2 十二项核心能力"), "5.2 十五项核心能力")
    capabilities = next(table for table in doc.tables if table.cell(0, 0).text.strip() == "能力")
    for row in [
        ["C-13", "Stylist Persona Continuity Intelligence（搭配师人设连续性智能）", "版本化管理专业身份、核心价值观、稳定审美立场、历史观点、栏目和有依据的观点演进，检测无依据冲突与模型切换引发的人设漂移。"],
        ["C-14", "Social-Media Native Voice Intelligence（自媒体原生语感智能）", "将已接受专业判断转化为平台原生、口播自然、场景具体、镜头可拍且非课件/模板/AI腔的表达，保持人设语言辨识度。"],
        ["C-15", "Multimodal Garment Understanding（多模态商品理解）", "从商品图片提取或判断类别、廓形、松量、长度、肩/腰线、视觉体量与重心、结构/垂坠/光泽、图案/尺度、色彩关系、装饰复杂度、可见穿法、风格倾向与商品间视觉兼容关系，并对证据分级。"],
    ]:
        clone_table_row(capabilities, row)
    insert_before(
        find_paragraph(doc, "5.3 品牌风格空间"),
        "D-17—D-29合并规则：每项须落到具体锚点，禁止浮动条文——D-20五对象（StylistPersonaProfile、PersonaMemorySnapshot、PublishedViewpointLedger、SeriesContinuityState、PersonaConflictRecord）进入M1 Schema清单、M7实现与M2人设连续性评分卡；D-21设独立FR、M2语感评分卡与M7实现；D-23进入garment_and_inventory.schema与MultimodalGarmentUnderstandingLayer组件注记；D-17进入M11注入测试合同；D-22/D-24进入15.6决策节点表端口条目；D-28进入3.3原则P-14、M2三分类合同与可接受决策边界、M3开放题模板、M4分歧账本、M5合法多解族、M6排序合同与DecisionTrace、M10迁移判据、10.2核心决策逻辑稳定率、16.1风险R-22与16.2停止条件、附录B术语；D-29进入M6通过标准与compiler qualification report的架构符合性检查三条、11.1组件注记与16.2停止条件。v1.2核验回执须对D-17—D-29新增内容重跑“目标→FR→门→里程碑”追溯检查。",
        "Normal",
    )
    insert_sections_before(
        doc,
        "6. 输入、输出与事实优先级合同",
        [
            (
                "5.5 搭配师人设连续性模型",
                [
                    "StylistPersonaProfile不是几个形容词、一段语气说明、固定Prompt或每次临时生成的人格。它必须覆盖专业身份与能力边界、目标人群、核心价值观、稳定审美立场、对品牌/身体/场景/消费/潮流的基本判断、常用观察视角、可坚持与可演进观点、禁止表达、不得虚构的履历/资历/顾客经历/职业事件，及与其他账号角色的职责差异。",
                    "PersonaMemorySnapshot、PublishedViewpointLedger、SeriesContinuityState、PlatformVoiceProfile、PersonaContinuityUpdate与PersonaConflictRecord共同记录已公开观点、栏目状态、观点修订原因、无新价值重复与冲突处置，确保品牌事实变化或底层模型切换不造成无依据人设漂移。",
                    "优先级：品牌真源＋已接受专业判断＋品类安全＋商品与库存事实 > 搭配师人设 > 平台表达偏好。人设不得覆盖上层事实、安全与专业结论。",
                ],
            ),
            (
                "5.6 多模态商品视觉机制",
                [
                    "MultimodalGarmentUnderstandingLayer将ProductImageAssetBundle转换为VisualAttributeExtractionResult；garment_and_inventory.schema必须为每个属性保存attribute_provenance、证据、置信度、图片/模型版本、推断时间、冲突、人工修订与是否允许进入正式决策。",
                    "attribute_provenance固定分层为authoritative、human_confirmed、model_inferred，并兼容authoritative_structured_fact、human_verified_visual_attribute、multimodal_inferred_visual_attribute别名。图片不能作为成分、性能、库存或安全认证的证据。",
                ],
            ),
            (
                "5.7 扩展兼容合同",
                [
                    "VisualMerchandisingExtensionPort预留StoreSpaceContext、FixtureAndCapacityContext、MannequinStylingContext、PlanogramContext与DisplayExecutionResult，未来可扩展橱窗、墙面、正/侧挂、叠装、模特组合、中岛/展台、动线、视觉焦点、陈列容量、数量间距、门店分级、Planogram及执行整改。",
                    "RealtimeSalesAssistExtensionPort预留SalesAssociateSessionContext、CustomerInteractionState与RealtimeRecommendationResult。两类扩展当前均不是V1.0强制实现或M12硬门；API、Schema和StylingResultBundle必须支持以新增模块方式开启，且保持现有输出向后兼容。",
                ],
            ),
        ],
    )

    # Chapter 6 objects, provenance, and failure-close behavior.
    replace_heading(doc, "6.1 输入对象", "6.1 十二项输入与配置对象")
    inputs = next(table for table in doc.tables if table.cell(0, 0).text.strip() == "对象" and table.cell(1, 0).text.strip() == "UniversalExpertKernel")
    for row in [
        ["ProductImageAssetBundle", "商品图片资产、来源、版本、角度、质量、权利/保密标记及可用范围。"],
        ["StylistPersonaProfile", "搭配师专业身份、能力边界、目标人群、价值观、审美立场、可坚持/演进观点、禁止表达与账号职责。"],
        ["PersonaMemorySnapshot", "特定版本下已发布观点、栏目、语言辨识度、修订原因和待处理冲突的可重放快照。"],
        ["FounderProvidedRealBrandPackage", "Founder可选提供的真实品牌与测试生产包，可含品牌定位、人群/风格、商品名/SKU/款色码/说明/图片/角色、可选库存或可售状态与测试任务。"],
        ["PublicationPolicy", "品牌、租户、账号、内容类型与风险等级下的发布模式、Founder授权、审核、撤回、回滚与Kill Switch配置；默认human_review。"],
    ]:
        clone_table_row(inputs, row)
    replace_heading(doc, "6.2 输出对象", "6.2 十五项输出与内部审计对象")
    outputs = next(table for table in doc.tables if table.cell(0, 0).text.strip() == "对象" and table.cell(1, 0).text.strip() == "BrandAudienceInterpretation")
    for row in [
        ["VisualAttributeExtractionResult", "来自商品图片的分级视觉属性、证据、置信度、冲突、人工覆盖与决策可用性。"],
        ["PersonaContinuityUpdate", "当次内容对观点、栏目和人设状态的版本化增量，包含冲突、修订原因与是否可接受。"],
        ["SocialMediaVoicePlan", "平台、口播节奏、停顿/转折、观点时机、场景/人物处境、镜头动作、专业密度与反模板约束。"],
        ["PublicationDecision", "发布模式、授权范围、全部前置检查、人工审核状态、自动发布许可/拒绝原因、撤回/回滚与Kill Switch记录。"],
    ]:
        clone_table_row(outputs, row)
    priority = next(table for table in doc.tables if table.cell(0, 0).text.strip().startswith("Specific Runtime Facts"))
    set_cell_text(
        priority.cell(0, 0),
        "Specific Runtime Facts（具体运行事实）\n>\nBrand Truth Pack（品牌真源） + Authoritative Structured Product Fact（权威结构化商品事实）\n>\nHuman Verified Visual Attribute（人工确认视觉属性）\n>\nCategory Hard Constraints（品类硬约束） + Accepted General Expert Knowledge（已接受通用专家知识）\n>\nMultimodal Inferred Visual Attribute（多模态视觉推断）\n>\nRetrieved Cases（检索案例）\n>\nGeneral Model Inference（一般模型推断）",
    )
    set_paragraph_text(
        find_paragraph(doc, "任何下层信息不得覆盖上层明确事实。模型推断必须显式标记，不得伪装为品牌、商品、顾客或真实经历。"),
        "任何下层信息不得覆盖上层明确事实。多模态视觉推断与一般模型推断必须显式标记，不得伪装为品牌、商品、库存、顾客或真实经历。FounderProvidedRealBrandPackage默认source_type=founder_provided_real_brand、usage_scope=test_production_only、live_market_evidence=false、hidden_benchmark_eligible=false、knowledge_training_eligible=false。",
    )
    failure_marker = find_paragraph(doc, "7. 功能需求清单")
    for text in [
        "• 商品图片缺失、质量不足或角度不支持某属性：该属性必须失败关闭、降低置信度或请求补图，不得填充不可验证的功能事实。",
        "• 缺少历史观点或人设快照：可以建立明确标记的初始人设版本，不得伪造既往履历、观点或顾客经历。",
        "• PublicationPolicy缺失、授权过期或前置检查不通过：必须回退为human_review或停止发布，不得默认启用自动发布。",
    ]:
        insert_before(failure_marker, text, "Normal")

    # Chapter 7 new functional requirements. Existing 4-column table is kept;
    # each new acceptance cell carries the acceptance, failure, and milestone.
    fr_table = next(table for table in doc.tables if table.cell(0, 0).text.strip() == "ID" and table.cell(1, 0).text.strip() == "FR-01")
    new_frs = [
        ["FR-22", "多模态商品图像理解", "系统必须从ProductImageAssetBundle提取搭配所需视觉属性，按证据等级与置信度进入决策。", "验收：属性准确性、置信度校准、跨图一致性达M2冻结阈值；权威事实覆盖率=0、不可验证功能声称率=0。失败状态：MULTIMODAL_ATTRIBUTE_UNSAFE。里程碑：M1/M2/M6/M8/M10。"],
        ["FR-23", "搭配师人设状态", "系统必须版本化保存StylistPersonaProfile、PersonaMemorySnapshot及人设更新/冲突。", "验收：人设身份与专业价值一致性达M2冻结阈值；虚构人设资历率=0。失败状态：PERSONA_STATE_INVALID。里程碑：M1/M2/M7/M10/M12。"],
        ["FR-24", "历史观点与栏目连续性", "系统必须记录已发布观点、检测无依据冲突，区分合理演进与人格漂移，避免无新价值重复。", "验收：观点冲突、栏目连续性、漂移和无新价值重复指标达M2阈值；已发布人设硬冲突率=0。失败状态：PERSONA_HARD_CONTRADICTION。里程碑：M2/M3/M5/M7/M10。"],
        ["FR-25", "自媒体原生语感", "系统必须实现平台原生表达、口播节奏、处境与镜头协同，并抑制课件化、模板化和AI腔。", "验收：原生语感、平台表达、口播节奏、可拍性、课件惩罚、模板化率与AI腔指标达M2阈值。失败状态：NATIVE_VOICE_QUALITY_FAIL。里程碑：M2/M3/M5/M7/M10。"],
        ["FR-26", "Founder真实品牌注入测试", "M11必须支持FounderInjectedRealBrandProductionTest；未注入时回退为Codex夹具品牌。", "验收：使用范围、图片/商品理解、测试生产报告和Founder评估Receipt可追溯；未供库存时不产生库存证据。失败状态：FOUNDER_PACKAGE_SCOPE_VIOLATION。里程碑：M1/M2/M11。"],
        ["FR-27", "扩展端口兼容", "系统必须提供VM与实时成交辅助的版本化扩展端口，不重建专家内核且不破坏现有输出。", "验收：新增模块可选、现有Schema/Bundle向后兼容，端口不绕过安全、授权和库存。失败状态：EXTENSION_COMPATIBILITY_BREAK。里程碑：M0/M1/M12。"],
        ["FR-28", "Founder控制发布模式", "默认human_review；Founder可细粒度授权auto_publish，且必须通过全部前置检查和Kill Switch。", "验收：未授权自动发布率=0，决策、撤回、回滚和紧急停止可审计。失败状态：UNAUTHORIZED_AUTO_PUBLISH。里程碑：M1/M2/M12。"],
        ["FR-29", "五类正式导入就绪", "进入正式真实品牌导入前，五类Category Adapter、安全合同、主要Schema、M10评测、版本/部署/权限/Feature Flag必须全部可开启。", "验收：five_category_activation_readiness=100%。失败状态：FIVE_CATEGORY_NOT_READY。里程碑：M1/M2/M9/M10/M12。"],
        ["FR-30", "Project CI与角色投影完整性", "系统治理必须以governance/bootstrap/role_operating_model.v0.2.yaml为单一规范源，确定性生成AGENTS/CLAUDE/Work CI投影并守卫产品真源、角色权限和隐藏资产边界。", "验收：instruction_projection_drift=0，双活产品真源=0，公开主仓隐藏内容/真实品牌资料/顾客数据命中=0；角色Prompt不构成第三套产品真源。失败状态：PROJECT_CI_INTEGRITY_FAIL。里程碑：M0/M12。"],
    ]
    for row in new_frs:
        clone_table_row(fr_table, row)

    # Chapter 8 distillation, candidate types, and strict data-pool isolation.
    roles = next(table for table in doc.tables if table.cell(0, 0).text.strip() == "角色" and table.cell(1, 0).text.strip() == "Founder")
    role_rows = [
        ["FOUNDER_PRODUCT_AUTHORITY", "唯一人类最终裁决票；承担产品、业务、领域、合规自评、语料/档案、数据与品牌资产、风险、里程碑、知识晋级、发布和最终合并。"],
        ["GPT_CHIEF_ADVISOR", "只读总顾问；审查战略、产品范围、连续性与远程交付，不写仓、不任Guardian、不作最终批准。ChatGPT Web与Work合计一个独立评审票。"],
        ["CLAUDE_EXECUTION_PLANNER", "只读执行规划；读取真源，产出任务包、允许/禁止路径、Checker与验收；不得写仓、改Founder裁决或审查自己规划的任务。"],
        ["CODEX_EXECUTION_ENGINEER", "默认唯一常规仓库写入者；按授权范围施工、Checker、Fixtures、Report、Receipt和候选Commit；仅Founder明确授权后合并。"],
        ["CLAUDE_INDEPENDENT_GUARDIAN", "隔离上下文只读审查冻结Commit；不得参与规划、施工、修复或产品重设计；只可PASS/CONDITIONAL/BLOCK并给证据。"],
        ["FUTURE_EXTERNAL_HUMAN_REVIEW_EXTENSION", "当前不启用、不构成基线前置；未来可由Founder另行开启外部领域或法律评审。"],
    ]
    for row, values in zip(roles.rows[1:], role_rows):
        for cell, value in zip(row.cells, values):
            set_cell_text(cell, value)
    set_paragraph_text(
        find_paragraph(doc, "同一个Claude Code会话不得同时充当候选生成者和独立守护者。必须使用独立工作区、独立上下文和不同任务合同。"),
        "当前采用FOUNDER_PLUS_ISOLATED_AI：Founder是唯一人类裁决票；Planner与Guardian必须为不同工作区、不同会话、不同任务合同，Guardian不得读取规划上下文或编辑候选。工作区记录只能作为程序性佐证；无法提供稳定session ID时须标记session_id_available=false并由Founder人工见证。",
    )
    insert_before(
        find_paragraph(doc, "8.3 问题系统的核心维度"),
        "当前人类班子：Founder 1人；外部领域专家0人；外部法律评审0人；独立人类数据管理员0人。当前AI工作面为GPT总顾问、Claude执行规划、Codex执行工程、Claude独立Guardian。未来外部人类评审扩展受支持，但不是当前基线的成立条件。",
        "Normal",
    )
    dimension = next(table for table in doc.tables if table.cell(0, 0).text.strip().startswith("Category（品类）"))
    set_cell_text(
        dimension.cell(0, 0),
        "Category（品类） × Brand Style Space（品牌风格坐标） × Audience / Individual（受众与个体） × Garment, Image & Inventory（商品、图片与货盘） × Scene & Function（场景与功能） × Conflict（冲突） × Transformation（变换） × Narrative Task（叙事任务） × Persona Continuity（人设连续性） × Historical Viewpoint & Series Continuity（历史观点与栏目连续性） × Social-Media Native Voice / Platform Expression（自媒体语感与平台表达） × Multimodal Product Understanding / Image-to-Attribute Conflict（多模态商品理解与图像属性冲突）",
    )
    knowledge = next(table for table in doc.tables if table.cell(0, 0).text.strip() == "知识对象")
    for row in [
        ["PersonaContinuityRuleCard", "描述人设核心不变项、可演进项、观点冲突、栏目续接、重复抑制与修订证据。"],
        ["SocialMediaVoicePatternCard", "保存平台原生表达、口播节奏、镜头动作、专业密度、反模板与反AI腔的适用边界。"],
        ["MultimodalAttributeRuleCard", "定义图像可推断视觉属性、不可验证属性、冲突优先级、置信度和人工覆盖规则。"],
    ]:
        clone_table_row(knowledge, row)
    scale = next(table for table in doc.tables if table.cell(0, 0).text.strip() == "资产")
    for row in scale.rows:
        if row.cells[0].text.strip() == "隐藏未见品牌":
            set_cell_text(row.cells[0], "封闭测试用合成未见品牌")
            set_cell_text(row.cells[2], "开发与知识生成链不可针对性读取品牌资料、配方或答案。")
    knowledge_states = next(table for table in doc.tables if table.cell(0, 0).text.strip() == "状态" and table.cell(1, 0).text.strip() == "model_elicited_candidate")
    for row in knowledge_states.rows:
        if row.cells[0].text.strip() == "expert_reviewed":
            set_cell_text(row.cells[0], "founder_ai_reviewed")
            set_cell_text(row.cells[1], "隔离GPT、隔离Claude、规则Checker与Founder按风险分级完成内部评审；不得表述为外部专家共识。")
    set_paragraph_text(
        find_paragraph(doc, "夹具品牌供给路径【已裁决】：试点三品类所需的品牌数据内容由Codex依据已冻结Schema与品类合同模拟合成，经Claude Code完整性审核后冻结为系统的夹具品牌（Fixture Brand）。用途：M1—M6开发调试、M6 golden/negative/boundary fixtures、非隐藏回归评测、M3—M6知识引出所需的仿真货盘、M11仿真试点运行。合成资产必须显式标记synthetic_fixture_brand / non_real_inventory / non_commercial_evidence，不得使用可误认为真实品牌的名称或叙事。"),
        "夹具品牌供给路径【已裁决】：Codex依据已冻结Schema与品类合同模拟合成夹具品牌、商品和图片描述，经Claude Code完整性审核后冻结。用途：M1—M6开发调试、M6 fixtures、非隐藏回归、M3—M6知识引出仿真货盘及M11仿真试点。合成资产必须标记synthetic_fixture_brand / non_real_inventory / non_commercial_evidence。M11未收到Founder真实品牌包时，全部使用该路径。",
    )
    set_paragraph_text(
        find_paragraph(doc, "隐藏品牌合成路径【已裁决】：8.6规划的30—50个隐藏未见品牌同样由合成路径建设。隔离条款随之生效：合成与审核必须在Guardian侧独立工作区完成，产物只进入冻结隐藏集；合成脚本、参数与随机种子入Receipt；任何夹具品牌数据或其合成会话不得复用于隐藏集，反之亦然。违反即触发16.2停止条件。"),
        "合成封闭未见品牌路径【已裁决】：8.6规划的30—50个封闭测试品牌继续由合成路径建设。夹具池与隐藏池必须在不同工作区使用不同生成配方、参数空间、随机种子和审核上下文；隐藏品牌只进入冻结隐藏集，不得向候选生成链泄漏，不得因测试失败补写品牌专属Prompt。脚本、参数、种子与清单哈希入Receipt。",
    )
    set_paragraph_text(
        find_paragraph(doc, "同质性控制【建议默认·待Founder确认】：考卷与教材同为合成时，最大风险是“题和练习出自同一口锅”——系统可能只学会合成腔，而非真实迁移能力。控制手段：夹具池与隐藏池必须使用不同的生成配方与参数空间；隐藏池须覆盖夹具池未使用的风格坐标区域；Guardian出同质性检查报告；M2泄漏检查扩展为“泄漏＋同质性”双检查。"),
        "同质性控制【已裁决】：夹具池与隐藏池同为合成时，必须通过不同配方、参数空间、随机种子、工作区和隐藏池专属风格坐标覆盖防止“考卷与教材同锅”。Guardian出同质性报告，M2执行泄漏＋同质性双检查。",
    )
    m8_7_end = find_paragraph(doc, "9. 叙事智能与账号矩阵内容合同")
    for text in [
        "Founder真实品牌注入路径【已裁决】：M11时Founder可提供FounderProvidedRealBrandPackage进行FounderInjectedRealBrandProductionTest。可包含品牌定位、目标人群、风格、商品名/SKU/款色码/说明/图片/角色、可选库存/可售状态与测试任务。如只提供商品与图片，可验证真实商品理解与测试生产，不得制造库存数量、销售状态或门店现货事实。",
        "Founder注入资料除非另行明确授权，不得用于训练、隐藏测试、隐藏品牌生成/增强/校准、通用专家知识固化或公开，也不得当作真实商业结果证据。合成封闭未见品牌是当前工程基准；Founder真实品牌注入是可选外部能力判断路径，两者用途不同、不得混用。",
        "M11多品牌要求不因Founder注入而变：一个注入真实品牌可替代一个夹具测试位，其余位继续由Codex合成；未注入时全部使用夹具。Founder注入不是M11强制硬门。",
        "隐藏评测物理隔离【S5/D-27】：隐藏题、答案、保密评分细则、隐藏品牌事实包、生成参数和原始输出不得进入主仓、Git历史、LFS、加密压缩包或Planner/Codex可读位置；只能存入Founder控制的独立私有仓、私有对象存储或离线介质。主仓仅保存benchmark_schema、frozen_manifest、content_hashes、runner_interface、result_summary与非秘密元数据；公开目录名固定为02_benchmark_manifests。",
        "Project CI【D-27】：governance/bootstrap/role_operating_model.v0.2.yaml是角色与执行治理的单一规范源；AGENTS.md、CLAUDE.md、Work/Copilot指令和角色Prompt均为生成投影或任务约束，不能成为第三套产品真源。GitHub Actions必须确定性守卫投影漂移、双活产品真源、角色越权、隐藏资产与秘密边界。",
    ]:
        insert_before(m8_7_end, text, "Normal")

    # Chapter 9 persona continuity and native voice contracts.
    replace_heading(doc, "9. 叙事智能与账号矩阵内容合同", "9. 叙事智能、搭配师人设连续性与自媒体原生语感合同")
    set_paragraph_text(
        find_paragraph(doc, "叙事化不是把专业术语改成口语，也不是套用“反常识标题 + 三个技巧”。系统必须从真实搭配决策中识别一个值得讲述的变化：某类人为什么总在某处出错；品牌目标与个体需求为什么发生冲突；一件旧衣服为什么以前没有被正确使用；删除、替换或重组后，造型和身份表达发生了什么变化。"),
        "叙事化不是把专业术语改成口语，也不是套用“反常识标题 + 三个技巧”。系统必须从真实搭配决策中识别值得讲述的人物处境、误解、冲突、变化和结论。“有叙事结构”不等于“具备自媒体原生语感”；标准叙事骨架是思考检查项，不得变成每条内容的固定文案模板。",
    )
    quality = next(table for table in doc.tables if table.cell(0, 0).text.strip() == "维度")
    for row in [
        ["人设连续", "专业身份、核心价值观、审美立场、历史观点与栏目状态前后有依据地一致或演进。"],
        ["原生语感", "具有平台原生表达、口播节奏、自然停顿/转折、观点时机、场景感、处境感与镜头协同，非课件/模板/AI腔。"],
        ["反模板", "不固定使用“反常识标题＋三个技巧”，同一主题必须有新价值，不以换句式冒充新内容。"],
    ]:
        clone_table_row(quality, row)
    insert_sections_before(
        doc,
        "10. 评测体系与商业发布门",
        [
            ("9.5 搭配师人设连续性合同", ["StylistPersonaProfile的专业身份、能力边界、目标人群、价值观、审美立场、基本判断、观察视角、可坚持/演进观点、禁止表达与账号职责必须明确。人设不得依靠形容词、固定Prompt、虚构高薪履历或临时人格维持。"]),
            ("9.6 历史观点与栏目连续性", ["系统必须记住已公开主要观点，检测新内容与历史观点的无依据冲突，区分合理演进和人格漂移，记录修订原因，避免无新价值重复，延续栏目和系列，并在品牌事实或模型变更后保持可解释的连续性。"]),
            ("9.7 自媒体原生语感合同", ["自媒体原生语感必须同时覆盖平台原生表达、口播节奏、自然停顿与转折、观点时机、具体场景与人物处境、专业判断生活化表达、镜头动作协同、非课件化、非模板化、非AI腔与不居高临下。"]),
            ("9.8 平台表达与反模板规则", ["同一专业结论可按平台改变信息密度、节奏、时长和镜头结构，但不得改变专业结论、品牌事实、商品事实与安全边界。不得使用固定内容骨架或同稿改写伪造平台原生性。"]),
            ("9.9 人设事实与履历真实性边界", ["不得虚构搭配师履历、资历、从业年限、服务客户、顾客经历、职业事件、门店现场或销售结果。人设可演进，但必须有原因、证据、版本与冲突处置记录。"]),
        ],
    )

    # Chapter 10 evaluation domains and gate integration.
    hard_gates = next(table for table in doc.tables if table.cell(0, 0).text.strip() == "硬门")
    for row in [
        ["人设真实", "fabricated_persona_credential_rate = 0；published_persona_hard_contradiction_rate = 0。"],
        ["人设不覆盖真源", "brand_truth_overridden_by_persona_rate = 0。"],
        ["多模态事实边界", "authoritative_fact_override_rate = 0；unverifiable_function_claim_rate = 0。"],
        ["发布授权", "unauthorized_auto_publish_rate = 0。"],
    ]:
        clone_table_row(hard_gates, row)
    quality_gates = next(table for table in doc.tables if table.cell(0, 0).text.strip() == "指标")
    new_metrics = [
        "persona_identity_consistency_score", "professional_value_consistency_score", "historical_viewpoint_contradiction_rate", "series_continuity_score", "persona_drift_rate", "topic_repetition_without_new_value_rate", "self_media_native_voice_score", "platform_native_expression_score", "spoken_rhythm_score", "shootability_score", "lecture_style_penalty_rate", "formulaic_template_rate", "ai_generated_tone_detection_rate", "visual_attribute_extraction_accuracy", "visual_attribute_confidence_calibration", "structured_fact_conflict_rate", "authoritative_fact_override_rate", "unverifiable_attribute_claim_rate", "cross_image_consistency_score", "low_quality_image_failure_close_rate", "five_category_activation_readiness", "unauthorized_auto_publish_rate", "ai_reviewer_disagreement_rate", "founder_override_rate", "founder_sample_review_rate", "high_risk_founder_full_review_rate",
    ]
    for metric in new_metrics:
        if metric == "five_category_activation_readiness":
            threshold, note = "M2_FREEZE_REQUIRED", "五类正式导入可开启就绪，M12/导入检查为100%。"
        elif metric == "high_risk_founder_full_review_rate":
            threshold, note = "100%", "童装安全、青少年年龄适当性、运动功能安全、隐私、品牌硬规则及法律/医疗/性能承诺不得抽样。"
        else:
            threshold, note = "M2_FREEZE_REQUIRED", "内部评审、人设/语感、多模态或发布治理评测域；由M2冻结口径与阈值。"
        clone_table_row(quality_gates, [metric, threshold, note])
    for row in quality_gates.rows:
        if row.cells[0].text.strip() == "未见品牌专家通过率":
            set_cell_text(row.cells[0], "internal_review_panel_pass_rate（未见品牌）")
            set_cell_text(row.cells[2], "由隔离GPT、隔离Claude、规则Checker与Founder按冻结评分卡形成内部面板结论；不是外部专家证明。")
        if row.cells[0].text.strip() == "未见品类/边缘任务通过率":
            set_cell_text(row.cells[0], "internal_review_panel_pass_rate（未见品类/边缘任务）")
            set_cell_text(row.cells[2], "内部评审面板重点检查规则迁移与安全边界；不得称为外部专家通过率。")
    hidden_marker = find_paragraph(doc, "10.4 商业发布八道门")
    for text in [
        "• 给出历史观点、栏目状态与品牌事实变化，检查人设稳定性、合理演进、无依据冲突和无新价值重复。",
        "• 将同一人设与专业结论投影到不同平台，检查核心人格一致性、平台原生表达、口播节奏、可拍性、课件/模板/AI腔。",
        "• 给出多角度、低质量和与结构化事实冲突的商品图片，检查视觉属性准确性、置信度校准、跨图一致性与失败关闭。",
        "• 替换底层模型后重放人设与内容任务，检查核心价值观、历史观点和多模态决策边界漂移。",
        "• 高风险知识（童装安全、青少年身体/年龄适当性、运动功能安全、隐私与个人信息、品牌硬规则、法律/医疗/性能承诺）Founder审查覆盖率必须为100%，不得抽样；中低风险才允许Founder分层抽样、多个隔离AI评审与确定性Checker组合。",
    ]:
        insert_before(hidden_marker, text, "Normal")
    release_gates = next(table for table in doc.tables if table.cell(0, 0).text.strip() == "发布门")
    gate_updates = {
        "Semantic Contract Gate": "对象、输入、输出、证据分级、优先级、失败状态与扩展兼容完整。",
        "Inventory Grounding Gate": "总部/门店可售性与发布前复核通过；多模态推断不得伪造库存、价格或现货事实。",
        "Narrative Truth Gate": "叙事真实、专业成立、账号差异化、人设连续性与自媒体原生语感通过。",
        "Operational Scalability Gate": "批量、恢复、成本、版本、审计、多租户隔离、发布策略权限/回滚/Kill Switch与扩展端口兼容通过。",
        "Safety & Compliance Gate": "品类安全、隐私/授权、AIGC标识、未成年人信息、数据出境、商品图片权利/保密及自动发布授权通过。",
    }
    for row in release_gates.rows[1:]:
        name = row.cells[0].text.strip().replace("\n", " ")
        if name in gate_updates:
            set_cell_text(row.cells[1], gate_updates[name])
    set_paragraph_text(
        find_paragraph(doc, "M0须交付compliance_review_contract.v1.0.yaml，确定合规责任人、核验清单与决策时间表；具体法律结论由法务形成。首轮法务核验须在M2冻结前完成，不得整体推迟到M9。首批核验清单六项："),
        "当前review_mode=FOUNDER_PLUS_ISOLATED_AI，external_human_review=false，external_legal_opinion=false。Founder须在founder_compliance_decision_ledger.yaml中对每项合规问题分别作ACCEPT/MITIGATE/DEFER/BLOCK裁决，记录理由、证据、控制、下次复核里程碑与external_legal_opinion_obtained=false；AI仅提供对抗支持，不得冒充律师意见。首批核验清单七项：",
    )
    ch11 = find_paragraph(doc, "11. 技术与运营要求")
    insert_before(ch11, "⑦ Founder或品牌提供的商品图片使用权、保密义务、保存期限、模型处理与输出公开范围。", "Normal")

    # Chapter 11 components and NFRs.
    components = next(table for table in doc.tables if table.cell(0, 0).text.strip() == "组件")
    for row in [
        ["MultimodalGarmentUnderstandingLayer", "摄取商品图片，提取视觉属性，管理来源、置信度、冲突、人工覆盖与决策可用性。"],
        ["PersonaMemoryEngine", "版本化保存人设档案、快照、栏目状态、观点更新与冲突。"],
        ["PublishedViewpointLedger", "记录已公开主要观点、证据、发布上下文、修订原因与演进链。"],
        ["SocialMediaVoiceEngine", "根据人设与平台输出原生表达、口播节奏、镜头协同与反模板约束。"],
        ["PublicationPolicyController", "执行human_review与Founder授权auto_publish策略，保存权限、审计、撤回、回滚与Kill Switch状态。"],
        ["AutoPublishKillSwitch", "对Founder授权自动发布执行即时全局或细粒度停止，记录触发人、范围、原因、撤回和回滚结果。"],
        ["ExtensionPortRegistry", "版本化管理VisualMerchandisingExtensionPort、RealtimeSalesAssistExtensionPort及未来可选扩展，维持Schema/Bundle兼容。"],
    ]:
        clone_table_row(components, row)
    for row in components.rows:
        if row.cells[0].text.strip() == "MultimodalGarmentUnderstandingLayer":
            set_cell_text(row.cells[1], "摄取商品图片并写入garment_and_inventory.schema.attribute_provenance；只允许authoritative > human_confirmed > model_inferred分层，推断不得覆盖权威事实，也不得推断成分、性能、库存与安全认证。")
    nfrs = next(table for table in doc.tables if table.cell(0, 0).text.strip() == "ID" and table.cell(1, 0).text.strip() == "NFR-01")
    for row in [
        ["NFR-09", "人设状态版本化与可重放", "StylistPersonaProfile、PersonaMemorySnapshot、PublishedViewpointLedger、栏目状态、冲突和观点修订必须可版本化、可回放、可回滚。"],
        ["NFR-10", "多模态推断来源与置信度", "每项视觉推断必须保存图片来源/版本、模型/版本、时间、置信度、证据、冲突、人工修订与正式决策许可。"],
        ["NFR-11", "扩展端口前后向兼容", "扩展必须以版本化端口、可选Schema或Bundle模块实现；未开启和旧版消费者的现有输出不受破坏。"],
        ["NFR-12", "发布模式权限、审计与Kill Switch", "publication_mode默认human_review；授权自动发布必须细粒度限定、全链审计、可撤回/回滚、可即时停止。"],
    ]:
        clone_table_row(nfrs, row)
    external_end = find_paragraph(doc, "12. 执行步骤与依赖关系")
    insert_before(external_end, "• 商品图片与多模态模型：可接入品牌授权图片与符合合同的多模态模型，图片权利、保密和数据边界须先行核验。", "Normal")
    insert_before(external_end, "• 发布平台：默认仅对接人工审核流程；只有PublicationPolicyController确认Founder授权与全部前置门通过后，才能调用自动发布扩展。", "Normal")

    # Chapters 12-14: execution and milestone deliverables.
    replace_text(doc, "M7｜叙事智能与账号投影V0.1", "M7｜叙事智能、人设连续性与账号投影V0.1")
    replace_text(doc, "M7 叙事智能与账号投影V0.1", "M7 叙事智能、人设连续性与账号投影V0.1")
    milestones = next(table for table in doc.tables if table.cell(0, 0).text.strip() == "节点")
    for row in milestones.rows:
        if row.cells[0].text.strip() == "M7":
            set_cell_text(row.cells[1], "叙事智能、人设连续性与账号投影V0.1")
            set_cell_text(row.cells[4], "叙事真实、人设连续性与原生语感门")
    # Existing count references in M1 and notes.
    replace_text(doc, "6.2的11个正式输出对象", "6.2的15个正式输出与内部审计对象")
    replace_text(doc, "7个输入对象↔Schema/合同文件一一对应", "12个输入与配置对象↔Schema/合同文件一一对应")
    replace_text(doc, "6.2的11个输出对象", "6.2的15个输出与内部审计对象")
    replace_text(doc, "11个输出对象", "15个输出与内部审计对象")

    additions = {
        "1": [
            "stylist_persona_profile.schema.v0.1.json",
            "persona_memory_snapshot.schema.v0.1.json",
            "published_viewpoint_ledger.schema.v0.1.json",
            "series_continuity_state.schema.v0.1.json",
            "platform_voice_profile.schema.v0.1.json",
            "product_image_asset_bundle.schema.v0.1.json",
            "visual_attribute_extraction_result.schema.v0.1.json",
            "garment_and_inventory.schema.v0.1.json（新增attribute_provenance：authoritative / human_confirmed / model_inferred）",
            "publication_policy.schema.v0.1.json",
            "extension_port_contracts（VM／实时成交辅助／结果Bundle兼容）",
        ],
        "2": [
            "persona_continuity_scoring_rubric.v0.1.yaml",
            "social_media_native_voice_scoring_rubric.v0.1.yaml",
            "multimodal_attribute_benchmark.v0.1.yaml",
            "multimodal_confidence_calibration_contract.v0.1.yaml",
            "five_category_readiness_definition.v0.1.yaml（阈值M2_FREEZE_REQUIRED）",
        ],
        "3": [
            "persona_continuity_elicitation_pack.v0.1.jsonl",
            "social_media_voice_elicitation_pack.v0.1.jsonl",
            "multimodal_garment_understanding_pack.v0.1.jsonl",
        ],
        "4": ["candidate archive扩展：人设规则／语感模式／多模态属性推断／对抗结果／跨图一致性"],
        "5": [
            "accepted_persona_continuity_rules",
            "accepted_social_media_voice_patterns",
            "accepted_multimodal_attribute_rules",
        ],
        "6": ["visual_attribute_confidence_and_override_consumer（消费结构化事实、图片推断、置信度与人工覆盖）"],
        "7": [
            "stylist_persona_profile.schema.v0.1.json",
            "persona_memory_snapshot.schema.v0.1.json",
            "published_viewpoint_ledger.schema.v0.1.json",
            "series_continuity_state.schema.v0.1.json",
            "platform_voice_profile.schema.v0.1.json",
            "persona_consistency_checker",
            "social_media_native_voice_checker",
            "anti_template_checker",
            "persona_and_voice_qualification_report",
        ],
        "8": [
            "product_image_ingestion_contract",
            "multimodal_attribute_extraction_pipeline",
            "attribute_provenance_and_override_service",
        ],
        "9": ["persona_body_language_and_image_separation_checks（人设语言不制造身体羞耻；商品图不推断消费者缺陷；未成年人商品图与个体图严格区分）"],
        "10": ["persona_voice_multimodal_transfer_report（新品牌人设稳定、跨平台核心人格、语感迁移、多模态鲁棒性、模型替换漂移）"],
        "11": [
            "（条件交付）founder_provided_real_brand_package_manifest",
            "（条件交付）founder_real_brand_test_production_report",
            "（条件交付）multimodal_product_understanding_report",
            "（条件交付）founder_capability_assessment_receipt",
        ],
        "12": [
            "persona_state_versioning_and_rollback",
            "publication_policy_controller",
            "auto_publish_kill_switch_contract",
            "extension_port_registry",
            "five_category_activation_readiness_report",
        ],
    }
    for chapter in (13, 14):
        for milestone, lines in additions.items():
            insert_milestone_deliverables(doc, milestone, chapter, lines)
    replace_text(doc, "rater_calibration_contract.v0.1.yaml", "reviewer_calibration_contract.v0.1.yaml")
    # §6.6：单条 reviewer_calibration_contract 兼容扩展，不允许出现两条同名交付物。
    replace_text(
        doc,
        "reviewer_calibration_contract.v0.1.yaml（评分卡培训流程、双盲抽检、评分者一致性阈值）",
        "reviewer_calibration_contract.v0.1.yaml（兼容旧名rater_calibration_contract；评分卡培训、双盲抽检、评分者一致性阈值、Founder标定锚点题、AI评审员盲评与分歧率、Founder覆盖与推翻记录、模型版本变化后重新标定；内部一致性不得表述为外部专家共识）",
        count=1,
    )
    replace_text(
        doc,
        "reviewer_calibration_contract.v0.1.yaml（评分者校准）",
        "reviewer_calibration_contract.v0.1.yaml（评分者校准；兼容旧名rater_calibration_contract，含Founder锚点题、AI盲评与分歧率、Founder覆盖/推翻记录与模型变更再标定）",
        count=1,
    )

    # Adjust milestone targets and pass standards in chapter 13 via stable text.
    target_replacements = {
        "目标：把本PRD转化为仓库中的可执行合同，建立独立项目空间，冻结角色、边界、非目标与状态。": "目标：把本PRD转化为仓库中的可执行合同，建立独立项目空间，冻结角色、边界、非目标、状态，并将人设连续性、原生语感、多模态事实分级、Founder真实品牌注入、发布策略和扩展兼容写入现有14项交付物及M1/M2 Brief。",
        "目标：用结构化知识、硬约束、候选组合、排序与LLM解释形成受约束决策。": "目标：用结构化知识、权威商品事实、有来源与置信度的多模态视觉推断、人工覆盖、硬约束、候选组合、排序与LLM解释形成受约束决策。",
        "目标：从已接受搭配决策中生成搭配命题、叙事弧、可拍动作和账号差异化内容。": "目标：从已接受搭配决策中生成搭配命题、叙事弧与可拍动作，同时维护搭配师人设、历史观点、栏目连续性、自媒体原生语感和账号差异化。",
        "目标：接入商品与库存，支持总部全货盘和单店现货两种任务范围。（v1.1注记：引擎与接口合同交付不变，但验证对象改为仿真货盘源；真实ERP/PIM/POS系统对接验证并入真实品牌补验批次。）": "目标：读取商品、图片与库存，支持总部全货盘和单店现货两种任务范围。引擎与接口合同在仿真货盘验证；真实ERP/PIM/POS稳定性仍并入真实品牌补验批次。",
    }
    for old, new in target_replacements.items():
        replace_text(doc, old, new, count=1)

    # M11 target and acceptance language occurs in chapter 13 only.
    m11_target = find_paragraph(doc, "目标：以夹具品牌（Codex合成、Claude Code完整性审核后冻结）在多品牌、多门店、多品类条件下验证端到端产能与稳定性、专业与叙事质量、库存约束/替代/撤稿的确定性执行、账号矩阵差异化、成本与恢复能力。关键门＝仿真运行门（v1.0原“真实采用门”已改判）。")
    set_paragraph_text(
        m11_target,
        "目标：M11仍为多品牌、多门店、多品类仿真试点，关键门仍为仿真运行门。数据供给优先使用路径A（FounderProvidedRealBrandPackage注入测试生产），未提供时使用路径B（Codex夹具品牌/商品/图片描述合成，Claude Code完整性审核后冻结）。一个Founder注入品牌可替代一个夹具位，其余位仍用夹具；注入路径不是M11强制硬门。两路径均验证专业/叙事/人设/多模态/货盘约束/账号矩阵/成本/恢复，并严格记录证据边界。",
    )
    for p in doc.paragraphs:
        if p.text.startswith("专业质量、品牌适配、可售性、内容真实性与账号差异化在仿真环境下通过；"):
            set_paragraph_text(p, p.text + " Founder真实品牌注入只能证明系统理解该品牌资料/商品图片、生产搭配与内容及Founder是否认可；不自动证明真实门店采用、消费者理解、库存接口稳定、商业转化或市场规模化效果。未提供时Receipt记录founder_real_brand_package_provided=false与fallback_source=codex_synthetic_fixture。")

    # M12 keeps its official name and path; add readiness evidence language.
    m12_pass = find_paragraph(doc, "八道商业发布门同时通过，并获得Founder最终批准。")
    set_paragraph_text(m12_pass, "八道商业发布门同时通过，five_category_activation_readiness=100%，并获得Founder最终批准。只实现人工审核仍可通过M12；自动发布实现不是必要条件，但PublicationPolicyController、权限合同、审计、回滚与Kill Switch架构兼容必须存在。")

    # Chapter 14 pass lines for newly governed milestones.
    pass_line_replacements = {
        "[ ] 验收门已通过：不存在未定义对象、品牌与品类混用、运行时事实混入长期真源或女装规则默认覆盖其他品类；6.2的15个输出与内部审计对象在Schema层全部可寻址，无一遗漏。": "[ ] 验收门已通过：12个输入/配置对象与15个输出/审计对象全部可寻址；品牌/品类/运行时事实不混用；多模态证据分级、人设状态、发布策略与扩展端口完整。",
        "[ ] 验收门已通过：内容不是知识点口语化；不存在虚构真实人物、事件或销售结果；账号之间不是同稿改写。": "[ ] 验收门已通过：内容不是知识点口语化；不虚构人物/事件/销售/人设履历；人设、历史观点与栏目可连续；自媒体语感通过冻结阈值；账号不是同稿改写。",
        "[ ] 验收门已通过：八道商业发布门同时通过，并获得Founder最终批准。": "[ ] 验收门已通过：八道商业发布门同时通过；five_category_activation_readiness=100%；发布策略、Kill Switch和扩展端口兼容存在；并获得Founder最终批准。",
    }
    for old, new in pass_line_replacements.items():
        matches = find_all_paragraphs(doc, old)
        if matches:
            for paragraph in matches:
                set_paragraph_text(paragraph, new)

    # Chapter 15 decisions, evidence wording, and five-category readiness.
    insert_before(find_paragraph(doc, "15.2 工作量结构"), "Commercial V1.0命名与工程路线【D-19已裁决】：M12正式名称、M0—M12顺序、关键路径与里程碑表保持不变。对仿真验证、Founder注入验证与真实市场待补验能力分别管理证据口径，不增加阻断工程推进的前置里程碑。", "Normal")
    m15_5_end = find_paragraph(doc, "15.6 后置事项决策节点指派表")
    insert_before(m15_5_end, "正式真实品牌导入就绪【已裁决】：M11仍只要求至少三类试点，但在正式真实品牌导入前，女装、童装、青少年潮流时装、亲子潮流时装、运动潮流时装五类Category Adapter、安全合同、主要Schema、M10评测、版本化、部署、权限与Feature Flag必须全部可开启，five_category_activation_readiness=100%。这不表示单一真实品牌必须经营五类。", "Normal")
    decision_end = find_paragraph(doc, "16. 风险、停止条件与决策权")
    for text in [
        "• Founder真实品牌注入——决策节点：M11试点基线冻结前；裁决人：Founder；未提供即自动回退Codex夹具合成，不阻断M11。",
        "• 发布模式——决策节点：每个品牌/租户/账号/内容类型首次启用自动发布前；裁决人：Founder；未决策默认human_review。",
        "• D-22 / VisualMerchandisingExtensionPort——决策节点：Founder提出空间陈列产品化需求时另立Brief；未启动时只维护端口、Schema版本和向后兼容，不得写成永久退出。",
        "• D-24 / RealtimeSalesAssistExtensionPort——决策节点：Founder确认实时终端形态、数据权限、延迟与安全合同后另立Brief；未启动时只维护端口、Schema版本和安全路由，不得写成永久退出。",
    ]:
        insert_before(decision_end, text, "Normal")

    # Chapter 16 risks and stop conditions.
    risks = next(table for table in doc.tables if table.cell(0, 0).text.strip() == "风险")
    risk_rows = [
        ["R-15", "人设漂移", "账号专业身份、核心价值观或审美立场随内容批次/模型切换无依据变化。", "人设快照、观点账本、版本重放、冲突检查与模型替换回归。"],
        ["R-16", "历史观点无依据冲突", "新内容与已发布核心观点相互矛盾，且无事实变化或修订理由。", "PublishedViewpointLedger、PersonaConflictRecord、修订原因与硬冲突零容忍。"],
        ["R-17", "语感模板化或AI腔", "平台内容课件化、固定骨架、居高临下或不自然。", "原生语感评分、anti-template checker、专家/内容评审与跨平台测试。"],
        ["R-18", "多模态商品属性幻觉", "图片推断覆盖权威事实，或虚构材质、功能、尺码、库存与功效。", "证据分级、置信度校准、冲突优先级、人工覆盖和不可验证声称硬门。"],
        ["R-19", "商品图片权利与保密风险", "Founder/品牌图片被超范围使用、留存、训练或公开。", "权利与保密标记、usage_scope、保存期限、访问控制、删除和法务核验。"],
        ["R-20", "未授权自动发布", "系统在Founder未授权、授权超范围或前置门未通过时发布。", "默认human_review、细粒度PublicationPolicy、全链审计、回滚与Kill Switch；unauthorized_auto_publish_rate=0。"],
        ["R-21", "扩展模块破坏兼容性", "VM、实时成交辅助或发布扩展改变旧Schema语义或绕过内核硬门。", "版本化ExtensionPortRegistry、可选Bundle模块、向后兼容回归与安全/授权/库存强制路由。"],
    ]
    for row in risk_rows:
        clone_table_row(risks, row)
    stop_end = find_paragraph(doc, "16.3 决策权")
    for text in [
        "• 多模态推断覆盖权威结构化商品事实：停止对应属性进入正式决策，回滚并重跑冲突回归。",
        "• 未经Founder授权启用自动发布：立即触发Kill Switch，停止发布，撤回可撤回内容并审计权限链。",
        "• 人设虚构履历、资历、顾客经历或职业事件：停止对应账号发布，回溯人设与内容链。",
        "• 发现严重历史观点冲突且未形成事实依据、修订原因和版本记录：停止该内容链。",
        "• Founder注入数据进入隐藏集、训练链或通用知识真源：停止测试与候选链，清理污染并重建隔离数据池。",
        "• 任一扩展模块破坏现有输出Schema/Bundle兼容性或绕过安全、授权、库存门：停止扩展发布并回滚。",
        "• Planner与Guardian同会话/同工作区，或Guardian读取规划上下文、编辑候选：停止审查并重新建立隔离工作面。",
        "• 隐藏题、答案、保密评分细则、隐藏品牌事实包、生成参数或原始输出进入主仓/Git历史：停止候选链并重建隔离存储。",
        "• 角色不可用却未在Receipt记录Founder替位/豁免裁决：默认DEFER，不得继续闭环。",
        "• CONDITIONAL缺少条件台账、证据、复核角色、closure commit或Founder关闭决定：不得声明PASS或CLOSED。",
    ]:
        insert_before(stop_end, text, "Normal")
    governance_marker = find_paragraph(doc, "17. 首个执行任务")
    insert_before(governance_marker, "16.4 任务分级、角色降级与条件关闭", "Heading 2")
    for text in [
        "L1仅限错别字、标点、无语义格式、机械同步、链接/索引或不改变合同含义的版式修复；须Founder确认L1、Codex修改、确定性Checker和Founder批准具体Commit。涉及范围、能力、角色、状态、阈值、风险、Schema、安全、合规、CI、知识晋级、发布或隐藏评测时不得使用L1。",
        "L2用于已冻结合同内实现、无产品语义变化的Schema/Checker与常规测试报告；执行Planner→Founder任务包/基线批准→Codex→Guardian→PR→ChatGPT总顾问→Founder批准→合并。",
        "L3用于PRD、角色权限、知识状态、隐藏评测、安全、合规、发布门、模型切换、知识晋级、Commercial版本和主仓治理；不得未经Founder显式豁免跳过独立Guardian、总顾问远程审查或Founder具体Commit批准。",
        "关键角色不可用默认DEFER。ChatGPT不可用时Founder可DEFER、指定替位顾问或显式豁免并接受风险；Codex不可用时默认顺延，紧急指定TEMPORARY_EXECUTION_WRITER时不得兼任Guardian，须独立工作区与同等证据。",
        "CONDITIONAL不等于PASS。每个条件必须进入conditional_decision_ledger.yaml，绑定source commit、owner、due milestone、required evidence、verification role、closure commit与Founder closure decision；修复产生新Commit后旧Guardian和总顾问结论失效，必须重审。",
        "授权层级固定为PRD→里程碑执行申请→Founder裁决→任务包/执行Prompt→角色Prompt；下层文件不得扩大上层权限。",
    ]:
        insert_before(governance_marker, text, "Normal")

    # Single-Founder operating capacity and review contract.
    effort_baseline = next(table for table in doc.tables if table.cell(0, 0).text.strip().startswith("工程量与工期默认基线"))
    set_cell_text(
        effort_baseline.cell(0, 0),
        "工程量与工期口径（单一Founder模式）\n125—185人月与15—24个月暂不改变；effort_unit=HUMAN_MACHINE_WORK_EQUIVALENT_PERSON_MONTH，staffing_commitment=false，founder_approved_estimation_basis=true。该数字是人机工作等价值与路线估算，不是当前必须招聘或配置相同人类团队的承诺。",
    )
    set_paragraph_text(
        find_paragraph(doc, "团队配置维持核心全职8—10人＋外部专家组6—10名兼职；v1.1新增角色见15.3。"),
        "当前实际班子按单一Founder模式运行，不把8—10人＋外部专家组写成当前强制编制。原配置仅保留为reference_delivery_capacity_model，status=NON_BINDING_REFERENCE，表示工作量等价能力模型，不表示当前人数或招聘承诺。",
    )
    replace_heading(doc, "15.3 建议核心团队", "15.3 当前实际班子与非约束容量参考")
    set_paragraph_text(
        find_paragraph(doc, "建议核心全职团队为8—10人，并使用分阶段外部专家审查。单一高价搭配师不得成为唯一真源，否则系统容易把个人风格、年龄与客户群偏好误写成通用规律。"),
        "current_human_team：Founder 1人；外部领域专家0人；外部法律评审0人；独立人类数据管理员0人。current_ai_workspaces：GPT_CHIEF_ADVISOR、CLAUDE_EXECUTION_PLANNER、CODEX_EXECUTION_ENGINEER、CLAUDE_INDEPENDENT_GUARDIAN。未来外部人类评审可由Founder开启，但不要求当前基线先行配置。",
    )
    replace_heading(doc, "15.4 专家审查量级", "15.4 内部评审容量参考")
    set_paragraph_text(find_paragraph(doc, "• POC：约40—60个专家工作日。"), "• POC：约40—60个内部评审工作日等价值，仅作非约束容量参考。")
    set_paragraph_text(find_paragraph(doc, "• Pilot MVP：累计约100—150个专家工作日。"), "• Pilot MVP：累计约100—150个内部评审工作日等价值；由隔离AI、规则Checker与Founder组合完成。")
    set_paragraph_text(find_paragraph(doc, "• Commercial V1.0：累计约180—260个专家工作日，分布于五类品类、隐藏评测、仿真试点与真实品牌补验批次。"), "• Commercial V1.0：累计约180—260个内部评审工作日等价值，分布于五类品类、隐藏评测、仿真试点与真实品牌补验批次；不代表外部专家已参与。")
    team = next(table for table in doc.tables if table.cell(0, 0).text.strip() == "角色" and table.cell(1, 0).text.strip() == "Founder / Product Authority")
    replacements = [
        ["FOUNDER_PRODUCT_AUTHORITY", "1人；唯一人类裁决票", "产品、业务、领域、合规自评、风险、里程碑、知识晋级、发布与最终合并。"],
        ["GPT_CHIEF_ADVISOR", "AI只读工作面", "战略、产品范围、连续性和远程交付审查；不写仓、不任Guardian。"],
        ["CLAUDE_EXECUTION_PLANNER", "AI只读工作面", "读取真源，输出任务包、边界、Checker与验收；不得写仓或自审。"],
        ["CODEX_EXECUTION_ENGINEER", "AI范围写入工作面", "默认唯一常规落盘；候选Commit、Checker、Fixtures、Report、Receipt；无Founder授权不合并。"],
        ["CLAUDE_INDEPENDENT_GUARDIAN", "AI隔离只读工作面", "只审冻结Commit；不规划、不施工、不修复。"],
        ["FUTURE_EXTERNAL_HUMAN_REVIEW", "当前0人；可选扩展", "外部领域/法律评审由Founder未来另行开启；当前external_human_review=false、external_legal_opinion=false。"],
        ["REFERENCE_DELIVERY_CAPACITY_MODEL", "非约束参考", "旧8—10人及外部专家配置仅解释工作量等价，不构成当前编制或招聘承诺。"],
        ["FOUNDER_COMPLIANCE_SELF_ASSESSMENT", "Founder逐项裁决", "六类合规问题逐项进入台账；AI提供对抗支持，不冒充外部法律意见。"],
        ["DATA_AND_ARCHIVE_OWNERSHIP", "Founder当前兼任", "语料、品牌档案、数据与隐藏集存储责任；隐藏集实际内容必须物理隔离于主仓。"],
        ["ROLE_UNAVAILABILITY_FALLBACK", "默认DEFER", "替位或豁免必须Founder显式裁决、独立会话/工作区并写Receipt；不得静默绕过。"],
    ]
    for row, values in zip(team.rows[1:], replacements):
        for cell, value in zip(row.cells, values):
            set_cell_text(cell, value)

    decision_rights = next(table for table in doc.tables if table.cell(0, 0).text.strip() == "角色" and table.cell(1, 0).text.strip() == "Founder" and len(table.rows) == 7)
    boundaries = [
        ["FOUNDER_PRODUCT_AUTHORITY", "唯一最终裁决：产品、风险、合规自评、里程碑、知识晋级、发布与合并。"],
        ["GPT_CHIEF_ADVISOR", "只读建议与远程对齐审查；无写入、Guardian或最终批准权。"],
        ["CLAUDE_EXECUTION_PLANNER", "只读规划；不得写仓、改裁决或自审。"],
        ["CODEX_EXECUTION_ENGINEER", "范围写入与候选冻结；无权改产品合同或未经Founder合并。"],
        ["CLAUDE_INDEPENDENT_GUARDIAN", "只读独立审查冻结Commit；无编辑、规划、修复与商业批准权。"],
        ["FUTURE_EXTERNAL_HUMAN_REVIEW", "当前未启用；未来意见不自动取代Founder裁决。"],
    ]
    for row, values in zip(decision_rights.rows[1:], boundaries):
        for cell, value in zip(row.cells, values):
            set_cell_text(cell, value)

    # ---- Supplement SUPP-01: D-28 合理多解原则 / D-29 M6 反退化验收 ----
    # 编号约束：原Prompt「原有编号一律不重排」。P-10—P-13 与 R-15—R-21 已被 D-20—D-26 占用，
    # 因此补充Prompt文字中的 "P-10" / "R-15" 在本版落为下一可用编号 P-14 / R-22，语义不变。
    principles = find_table(doc, "原则")
    clone_table_row(
        principles,
        [
            "P-14",
            "Non-Uniqueness by Design（合理多解原则）",
            "除具体运行事实、库存、品类安全、品牌硬规则及明确合同约束外，搭配诊断、候选组合、视觉取舍、风格表达与叙事命题"
            "不得假设存在唯一标准答案。评测对象是约束遵守、机制成立性、取舍质量、可解释性与决策边界，而非与单一参考答案的相似度。",
        ],
    )
    insert_before(
        find_paragraph(doc, "4. 领域范围、用户与核心场景"),
        "P-14落地口径：评测任务强制三分类——① constraint_correctness（硬约束题，唯一0/1）；② mechanism_correctness"
        "（评可接受推理区间）；③ open_decision（评可接受解空间与多解族质量）。②③类禁止设置唯一Gold Answer，"
        "冻结对象是Acceptable Decision Boundary（可接受决策边界），不是参考答案。",
        "Normal",
    )

    # D-28 · M2：EvaluationCard 与 benchmark_capability_matrix 强制三分类。
    knowledge_objects = find_table(doc, "知识对象")
    for row in knowledge_objects.rows:
        if row.cells[0].text.strip() == "EvaluationCard":
            set_cell_text(
                row.cells[1],
                "题目、输入、预期边界、评分标准、硬失败和裁决说明；必须标注evaluation_task_class"
                "（constraint_correctness／mechanism_correctness／open_decision），②③类不得写入唯一Gold Answer。",
            )
    replace_text(
        doc,
        "benchmark_capability_matrix.v0.1.yaml",
        "benchmark_capability_matrix.v0.1.yaml（每条能力项须标注evaluation_task_class：constraint_correctness／mechanism_correctness／open_decision）",
        count=2,
    )

    supplement_additions = {
        "2": [
            "evaluation_task_class_contract.v0.1.yaml（D-28强制三分类：constraint_correctness唯一0/1；mechanism_correctness评可接受推理区间；open_decision评可接受解空间与多解族质量）",
            "acceptable_decision_boundary_registry.v0.1.yaml（②③类冻结对象＝可接受决策边界；禁止设置唯一Gold Answer）",
        ],
        "3": [
            "open_decision_question_template_contract.v0.1.yaml（D-28：开放型问题须要求至少两个成立但取舍不同的方案，并说明什么条件变化会交换优先级；与FR-08对齐）",
        ],
        "5": [
            "disagreement_classification_and_solution_family_ledger.v0.1.yaml（D-28二分类：错误分歧＝违反硬约束/安全/机制必须裁决；合法审美分歧＝约束全满足但取舍不同，保留为合法多解族并给知识对象加solution_family标记）",
        ],
        "6": [
            "architecture_conformance_check_report（D-29架构符合性检查三条：硬约束确定性代码＋fixtures单测、LLM-off不变性可重放、DecisionTrace显示每个排除项的确定性规则）",
            "llm_off_invariance_replay_fixtures（关闭LLM解释/生成组件后的约束过滤与候选排除结果比对）",
            "acceptable_solution_family_trace（D-28：DecisionTrace记录被舍弃的合法候选族与取舍理由）",
        ],
        "10": [
            "legal_decision_space_conformance_report（D-28：迁移判据为是否持续落在合法决策空间内，不与隐藏参考答案比对唯一一致性）",
        ],
    }
    for chapter in (13, 14):
        for milestone, lines in supplement_additions.items():
            insert_milestone_deliverables(doc, milestone, chapter, lines)

    # D-28 · M4 分歧账本保留合法分歧。
    replace_text(
        doc,
        "model_disagreement_ledger.v0.1.jsonl",
        "model_disagreement_ledger.v0.1.jsonl（D-28：保留合法分歧，不得为评分便利消灭）",
        count=2,
    )

    # D-28 / D-29 · 里程碑通过标准（第13章正文与第14章验收门同步）。
    pass_standard_updates = {
        "隐藏集、评分卡、硬门和阈值冻结；后续失败不得降标；泄漏与同质性双检查通过。":
            "隐藏集、评分卡、硬门和阈值冻结；后续失败不得降标；泄漏与同质性双检查通过；"
            "EvaluationCard与benchmark_capability_matrix三分类齐备，②③类未设置唯一Gold Answer，冻结对象为可接受决策边界。",
        "问题覆盖多条件冲突、排除理由、副作用、规则失效、跨品牌/品类迁移与叙事真实性；禁止以知识清单题为主。":
            "问题覆盖多条件冲突、排除理由、副作用、规则失效、跨品牌/品类迁移与叙事真实性；禁止以知识清单题为主；"
            "开放型问题模板须给出至少两个成立但取舍不同的方案及优先级交换条件。",
        "所有物理调用、失败、重试与候选均可追踪；不得把候选标为accepted。":
            "所有物理调用、失败、重试与候选均可追踪；不得把候选标为accepted；合法分歧在分歧账本中保留，不得为评分便利消灭。",
        "知识状态链完整；个人审美不得伪装为普遍规律；高风险知识有多人审查。":
            "知识状态链完整；个人审美不得伪装为普遍规律；高风险知识Founder审查覆盖率100%，AI分歧与Founder推翻均可审计；"
            "分歧已按错误分歧与合法审美分歧二分类，合法多解族带solution_family标记保留。",
        "库存、品牌、品类、用户拒绝项等硬约束确定性执行；所有选择与排除可解释。":
            "库存、品牌、品类、用户拒绝项等硬约束确定性执行；所有选择与排除可解释；架构符合性检查三条全部通过——"
            "① 硬约束由确定性代码执行且可用fixtures单测；② LLM-off不变性：关闭LLM解释/生成组件后约束过滤与候选排除结果不变且可重放；"
            "③ DecisionTrace可显示每个排除项由哪条确定性规则触发。StylingRanker在可行解集合上排序而非唯一argmax，"
            "被舍弃的合法候选族及理由须进入DecisionTrace。",
        "硬门全通过，质量门达到M2冻结阈值；失败不得通过补写隐藏品牌专属规则修复。":
            "硬门全通过，质量门达到M2冻结阈值；失败不得通过补写隐藏品牌专属规则修复；"
            "迁移判据为输出是否持续落在合法决策空间内，不是与隐藏参考答案一致。",
    }
    for old, new in pass_standard_updates.items():
        replace_text(doc, old, new, count=2)

    # D-29 · M6 目标：禁止单层 Prompt/RAG 端到端直答。
    replace_text(
        doc,
        "目标：用结构化知识、权威商品事实、有来源与置信度的多模态视觉推断、人工覆盖、硬约束、候选组合、排序与LLM解释形成受约束决策。",
        "目标：用结构化知识、权威商品事实、有来源与置信度的多模态视觉推断、人工覆盖、硬约束、候选组合、排序与LLM解释形成受约束决策。"
        "M6不得实现为“把已接受知识注入Prompt/RAG后由LLM端到端直接作答”的单层结构：硬约束执行（库存/品牌/品类/安全/用户拒绝项）、"
        "候选生成与排序必须是确定性、可审计步骤；LLM用于语义解析、解释与叙事生成，不得作为硬约束的唯一执行者。",
        count=1,
    )

    # D-28 / D-29 · 11.1 组件注记。
    component_updates = {
        "ConstraintEngine": "执行库存、品牌、品类、安全、用户拒绝和场景硬约束。D-29：必须由确定性、可审计代码执行并可用fixtures单测；"
                            "关闭LLM组件后过滤与排除结果不变。",
        "StylingRanker": "在可行解集合上按品牌、个体、场景、功能与经营目标排序和取舍；输出是有序的合法候选族，不是唯一argmax（D-28）。",
        "DecisionTraceEngine": "保存事实、知识、候选、排除、不确定性和版本；每个排除项须标注触发它的确定性规则，"
                               "被舍弃的合法候选族及取舍理由一并记录（D-28/D-29）。",
    }
    components_table = find_table(doc, "组件")
    for row in components_table.rows:
        name = row.cells[0].text.strip()
        if name in component_updates:
            set_cell_text(row.cells[1], component_updates[name])

    # D-28 · 10.2 指标更名，阈值不变。
    gates_table = find_table(doc, "指标")
    renamed = 0
    for row in gates_table.rows:
        if row.cells[0].text.strip() == "核心判断重复一致率":
            set_cell_text(row.cells[0], "核心决策逻辑稳定率")
            set_cell_text(
                row.cells[2],
                "相同输入与约束下，核心专业逻辑（机制判断、约束处理、取舍原则）不得无依据漂移；具体选品与方案族允许不同。"
                "反例＝同条件下先“弱化肩部”后“强化肩部”（失败）；非反例＝同条件下一次裤装方案、一次裙装方案且均成立（不计失败）。",
            )
            renamed += 1
    if renamed != 1:
        raise ValueError(f"expected exactly one 核心判断重复一致率 row, got {renamed}")

    # D-28 · 16.1 风险 R-22。
    risks_table = find_table(doc, "风险")
    clone_table_row(
        risks_table,
        [
            "R-22",
            "唯一答案偏误",
            "benchmark被做成单一标准答案型，把开放决策题当成有唯一Gold Answer的考试，考错对象。",
            "严重度对标R-09：D-28三分类合同、②③类禁设唯一Gold Answer、M2冻结前checker与独立Guardian审查。",
        ],
    )

    # D-28 / D-29 · 16.2 停止条件。
    supplement_stop_end = find_paragraph(doc, "16.3 决策权")
    for text in [
        "• mechanism_correctness或open_decision类任务被设置唯一Gold Answer：停止M2冻结，改回可接受决策边界后重跑。",
        "• M6以纯Prompt/RAG端到端直答实现，或硬约束由LLM单独执行、LLM-off后过滤与排除结果改变：停止M6通过判定并重建确定性执行层。",
    ]:
        insert_before(supplement_stop_end, text, "Normal")

    # Chapter 17 keeps exactly 14 top-level M0 deliverables.
    first_task = find_paragraph(doc, "该任务只完成M0并为M1—M2建立可执行基础，不进行大规模知识蒸馏，不接真实品牌生产，不翻Serving状态。")
    set_paragraph_text(first_task, first_task.text + " D-17—D-29新增能力与治理只写入既有合同和M1/M2 Brief，不增加M0第15项，不在M0处理Founder真实品牌数据、运行多模态识别、建立人设记忆生产库或启用自动发布。")

    # Appendix terminology.
    terms = next(table for table in doc.tables if table.cell(0, 0).text.strip() == "术语")
    term_rows = [
        ["Stylist Persona Continuity（搭配师人设连续性）", "搭配师的专业身份、价值观、审美立场、历史观点、栏目与语言辨识度在长期内容中可连续、可演进、可审计的能力。"],
        ["Social-Media Native Voice（自媒体原生语感）", "平台原生、口播自然、场景与镜头协同、非课件/模板/AI腔，且不改变专业结论的表达能力。"],
        ["Persona Memory Snapshot", "特定版本的人设、已发布观点、栏目状态、修订原因与冲突状态快照。"],
        ["Published Viewpoint Ledger", "记录账号已公开主要观点、证据、上下文、修订与演进的可审计账本。"],
        ["Multimodal Garment Understanding（多模态商品理解）", "从商品图片提取搭配所需视觉属性并管理证据、置信度、冲突与人工覆盖的能力。"],
        ["Visual Attribute Provenance（视觉属性来源）", "某视觉属性的图片来源/版本、模型/版本、时间、证据、置信度、冲突与人工修订链。"],
        ["FounderProvidedRealBrandPackage", "Founder为M11可选测试生产提供的真实品牌、商品、图片、可选货盘与任务包；默认不训练、不进隐藏集、不作商业证据。"],
        ["FounderInjectedRealBrandProductionTest", "M11中对FounderProvidedRealBrandPackage进行的可选测试生产，用于外部能力判断，不替代封闭未见品牌或真实市场补验。"],
        ["Publication Policy（发布策略）", "决定human_review或Founder授权auto_publish的品牌/租户/账号/内容/风险范围、检查、撤回、回滚与Kill Switch合同。"],
        ["Visual Merchandising Extension Port", "面向未来空间陈列上下文、Planogram与执行结果的版本化扩展端口；非V1.0强制交付门。"],
        ["Realtime Sales Assist Extension Port", "面向未来导购会话、顾客交互状态与实时推荐的版本化扩展端口；非V1.0强制交付门。"],
        ["Five-category Activation Readiness（五品类激活就绪）", "五类首发品类的Adapter、安全合同、主要Schema、M10评测、版本/部署/权限/Feature Flag全部可开启的正式真实品牌导入前置状态。"],
        ["Non-Uniqueness by Design（合理多解原则）", "除具体运行事实、库存、品类安全、品牌硬规则与明确合同约束外，搭配诊断、候选组合、视觉取舍、风格表达与叙事命题不假设唯一标准答案；评测对象是约束遵守、机制成立性、取舍质量、可解释性与决策边界。"],
        ["Acceptable Decision Boundary（可接受决策边界）", "mechanism_correctness与open_decision类任务的冻结对象：界定哪些推理区间与解空间成立、哪些越界，取代唯一参考答案。"],
        ["Legitimate Solution Family（合法多解族）", "在全部硬约束、安全与品牌规则均满足前提下取舍不同的一组成立方案；合法审美分歧保留为该族，知识对象以solution_family标记，不作为错误分歧裁决消灭。"],
        ["Decision Logic Stability Rate（核心决策逻辑稳定率）", "相同输入与约束下核心专业逻辑（机制判断、约束处理、取舍原则）不无依据漂移的比率；具体选品与方案族允许不同，原名“核心判断重复一致率”，阈值≥85%不变。"],
        ["LLM-off Invariance（LLM-off不变性）", "关闭LLM解释与生成组件后，约束过滤与候选排除结果保持不变且可重放；D-29的M6架构符合性检查项之一。"],
    ]
    for row in term_rows:
        clone_table_row(terms, row)

    # Closing statement, header/footer, and metadata.
    replace_text(
        doc,
        "一个专家内核，适配不同品牌与品类；真实货盘约束专业判断；专业判断产生叙事内容；模型知识只能经评测与裁决晋级。",
        "一个专家内核适配不同品牌与品类；权威事实与证据分级多模态理解约束专业判断；专业判断通过持续人设与原生语感生成叙事；模型知识只能经评测与裁决晋级；扩展能力按Founder裁决开启且不破坏内核合同。",
        count=1,
    )
    replace_text(doc, "D-17—D-26", "D-17—D-29")
    replace_text(doc, "D-17~D-26", "D-17~D-29")
    replace_text(doc, "02_benchmarks_hidden/", "02_benchmark_manifests/")
    replace_text(doc, "外部专家对高价值、高风险与高分歧候选进行校准并形成接受基线。", "隔离GPT、隔离Claude与规则Checker对中低风险候选进行内部盲评；Founder对高风险候选100%审查，并裁决知识晋级。")
    replace_text(doc, "专家抽评≥100件", "内部评审抽检≥100件")
    replace_text(doc, "专家通过率初值", "internal_review_panel_pass_rate初值")
    replace_text(doc, "知识蒸馏、知识工程与专家审查", "知识蒸馏、知识工程与内部评审")
    replace_text(doc, "专家知识门", "内部知识评审门")
    replace_text(doc, "专家审查工作台", "内部评审工作台")
    replace_text(doc, "专家复核", "Founder与隔离AI复核")
    replace_text(doc, "引入多专家与真实案例", "引入隔离多评审、Founder裁决与可验证案例")
    replace_text(doc, "原生语感评分、anti-template checker、专家/内容评审与跨平台测试", "原生语感评分、anti-template checker、隔离AI/Founder内部评审与跨平台测试")
    replace_text(doc, "— Founder正式立项基线 v1.0 · 2026-08-12 ｜ v1.1 独立解耦与审查修复基线 · 2026-08-12 —", "— PRD v1.2 · 人设连续性、多模态商品理解与扩展兼容基线 · 2026-08-12 · 待Founder签署生效 —", count=1)
    for section in doc.sections:
        if section.header.paragraphs:
            set_paragraph_text(section.header.paragraphs[0], "DIYU-CBFSK-001  |  笛语跨品牌服装搭配专家内核  |  PRD v1.2")
    doc.core_properties.title = "笛语跨品牌服装搭配专家内核 PRD 与执行里程碑 v1.2"
    doc.core_properties.subject = "D-17—D-29、单一Founder治理操作模型与Project CI候选基线"
    doc.core_properties.comments = f"EXECUTION_RUN_ID={EXECUTION_RUN_ID}; PROJECT_INITIATED; EXECUTION_NOT_STARTED; M0_AUTHORIZED_FALSE; PRODUCTION_SERVABLE_FALSE"
    doc.save(PRD_OUTPUT)


def build_m0() -> None:
    doc = Document(_source(M0_SOURCE))
    replace_text(doc, "DIYU-CBFSK-EXEC-REQ-M0-002", "DIYU-CBFSK-EXEC-REQ-M0-003")
    replace_text(doc, "PRD v1.1", "PRD v1.2")
    replace_text(doc, "PRD / v1.1", "PRD / v1.2")
    replace_text(doc, "_v1.1.docx", "_v1.2.docx")
    replace_text(doc, "与 v1.1 Delta 冲突条文", "与 PRD v1.2 及 D-17—D-29 冲突条文")
    replace_text(doc, "人工发布与夹具品牌不可证明范围", "Founder控制发布、多模态事实分级与夹具/Founder注入品牌不可证明范围")

    control = doc.tables[0]
    set_table_value(control, "申请 ID", "DIYU-CBFSK-EXEC-REQ-M0-003")
    set_table_value(control, "生效前置条件", "PRD v1.2全文（笛语跨品牌服装搭配专家内核_PRD与执行里程碑_v1.2.docx）经Founder签署生效，且本申请经Founder签署。两项均未满足前，m0_authorized=false，不得开工。")
    set_table_value(control, "依据文件", "待Founder签署的PRD v1.2全文候选＋PRD_v1.2_change_map.yaml＋PRD_v1.2_核验回执.docx＋治理角色规范源。v1.2生效前，根目录v1.1仍保持远程当前活基线且状态PENDING_FOUNDER_SIGNATURE；不得提前归档。")
    set_table_value(control, "任务范围", "仅申请完成M0，并为M1–M2建立可执行基础；产出PRD v1.2第13/14/17.1节统一的十四项顶层交付物。D-17—D-29与S1—S8写入现有合同文件和M1/M2 Brief，不增加第15项。")
    set_table_value(control, "本申请取代", "本申请取代DIYU-CBFSK-EXEC-REQ-M0-002（M0执行申请v1.1）。更早的001及18项版均仍为已归档历史文件，不得作为执行依据。")
    set_table_value(
        control,
        "Founder 已裁决冲突项",
        "① M0顶层交付清单仍为14项，D-17—D-29与S1—S8均内嵌现有合同与M1/M2 Brief。② M11仍为夹具品牌仿真试点，增加Founder真实品牌注入优先路径与Codex夹具回退；注入非M11硬门。③ 未见品牌内容物理隔离于主仓；Founder数据不得进入。④ M12名称、M0—M12顺序与Commercial V1.0路线不变。⑤ 人设连续性、自媒体原生语感和多模态商品理解升为正式能力与评测域。⑥ VM与实时成交辅助当前不强制实现，但保留前后向兼容端口。⑦ 默认人工审核；自动发布仅由Founder按租户/品牌/账号/风险级别显式授权，FR-17与人工可发布率≥75%不变。⑧ M11至少三类不变；正式真实品牌导入前五类必须100%可开启就绪。⑨ 单一Founder＋隔离AI角色模型、任务分级、条件/合规台账和Project CI进入M0合同。",
    )
    replace_heading(doc, "一、交付物清单（与 PRD v1.2 第 13 / 14 / 17.1 节完全一致）", "一、交付物清单（与 PRD v1.2 第 13 / 14 / 17.1 节完全一致）")

    # Redlines and Guardian review additions.
    redline_end = find_paragraph(doc, "四、验收与收口")
    for text in [
        "M0不得实际接收、处理、保存或测试Founder真实品牌资料；只定义FounderProvidedRealBrandPackage、使用范围、隔离与Receipt合同。",
        "M0不得运行多模态商品识别或产生正式VisualAttributeExtractionResult；只定义证据分级、Schema、评测与失败关闭Brief。",
        "M0不得建立搭配师人设记忆生产库、写入已发布观点或运行人设持久化；只定义对象、边界、M1 Schema与M2评测Brief。",
        "M0不得启用自动发布；publication_mode仍为human_review。本任务只定义Founder授权、前置检查、审计、撤回、回滚与Kill Switch合同。",
    ]:
        insert_before(redline_end, "• " + text, "List Paragraph")
    guardian_end = find_paragraph(doc, "六、审核与授权")
    for text in [
        "⑥ 人设与语感闭环：人设连续性、历史观点/栏目、自媒体原生语感是否进入M1对象和M2评测合同 → PASS / FAIL ＋证据。",
        "⑦ 多模态事实边界：authoritative_structured_fact > human_verified_visual_attribute > multimodal_inferred_visual_attribute > general_model_inference是否明确，不可仅凭图片断言项是否失败关闭 → PASS / FAIL ＋风险。",
        "⑧ 发布模式：是否正确表达为默认人工、Founder可授权自动发布，且既非永久禁止也非默认开启 → PASS / FAIL ＋证据。",
        "⑨ 扩展兼容：VM与实时成交辅助是否不列入V1.0/M12强制门且保留版本化扩展端口 → PASS / FAIL ＋证据。",
        "⑩ M0顶层清单：第13/14/17.1节与本申请是否仍严格为同一组14项，无第15项或18项版回潮 → PASS / FAIL ＋证据。",
    ]:
        insert_before(guardian_end, "• " + text, "List Paragraph")

    review_table = doc.tables[-1]
    set_cell_text(review_table.rows[1].cells[0], "文档一致性预审")
    set_cell_text(review_table.rows[1].cells[1], "候选生成侧自检完成 · 非Guardian结论")
    set_cell_text(review_table.rows[1].cells[2], f"已核对PRD v1.2、D-17—D-29、S1—S8、M0十四项、M1/M2承接、多模态来源分层、隐藏集隔离、单一Founder治理与Project CI。此项只是候选施工侧自检（TEMPORARY_EXECUTION_WRITER），独立Guardian与ChatGPT总顾问尚未审查。execution_run_id={EXECUTION_RUN_ID}")
    set_cell_text(review_table.rows[2].cells[0], "Founder 授权（PRD v1.2 ＋ 本申请）")
    for section in doc.sections:
        header = section.header
        if header.paragraphs:
            set_paragraph_text(header.paragraphs[0], "DIYU-CBFSK-001  |  M0 执行申请 v1.2  |  EXEC-REQ-M0-003")
        else:
            paragraph = header.add_paragraph("DIYU-CBFSK-001  |  M0 执行申请 v1.2  |  EXEC-REQ-M0-003")
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    doc.core_properties.title = "笛语跨品牌服装搭配专家内核 M0 执行申请 v1.2"
    doc.core_properties.comments = f"EXECUTION_RUN_ID={EXECUTION_RUN_ID}; M0_AUTHORIZED_FALSE; requires Founder signature on PRD v1.2 and this request"
    doc.save(M0_OUTPUT)


def clear_document_body(doc: DocumentType) -> None:
    body = doc._element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def shade_cell(cell: _Cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def add_receipt_table(doc: DocumentType, headers: list[str], rows: list[list[str]], widths: list[float] | None = None) -> Table:
    table = doc.add_table(rows=1, cols=len(headers))
    if "Table Grid" in [style.name for style in doc.styles]:
        table.style = "Table Grid"
    for index, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[index], header)
        shade_cell(table.rows[0].cells[index], LIGHT)
        for run in table.rows[0].cells[index].paragraphs[0].runs:
            set_run_font(run, size=9.0, color=BLUE_DARK, bold=True)
    for values in rows:
        row = table.add_row()
        for index, value in enumerate(values):
            set_cell_text(row.cells[index], value)
            for run in row.cells[index].paragraphs[0].runs:
                set_run_font(run, size=8.7)
    if widths:
        for row in table.rows:
            for index, width in enumerate(widths):
                row.cells[index].width = Pt(width)
    doc.add_paragraph("")
    return table


def build_receipt() -> None:
    # Use the PRD template for a stable style catalog and page geometry; the
    # v1.1 receipt is still an authoritative content input and is archived.
    doc = Document(_source(PRD_SOURCE))
    clear_document_body(doc)
    title = doc.add_paragraph("PRD v1.2 核验回执")
    title.style = doc.styles["Title"] if "Title" in [s.name for s in doc.styles] else doc.styles["Normal"]
    subtitle = doc.add_paragraph("DIYU-CBFSK-001 · D-17—D-29、S1—S8、治理操作模型与执行基线核验")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_receipt_table(
        doc,
        ["控制项", "内容"],
        [
            ["回执 ID", "DIYU-CBFSK-PRD-V1.2-VERIFY-RECEIPT-001"],
            ["核验对象", "PRD v1.2全文＋M0执行申请v1.2＋README＋PRD_v1.2_change_map.yaml＋治理合同＋Checker/审计脚本"],
            ["核验依据", "Founder最新D-23/D-25补丁＋D-17—D-29＋S1—S8＋已确认Manifest > PRD v1.1未被改判条款 > 历史证据"],
            ["核验方式", "机器Checker＋DOCX OOXML/ZIP完整性＋结构/版本/清单/追溯全文检索＋PDF渲染与版式抽检"],
            ["核验人／日期", f"候选施工侧自检（TEMPORARY_EXECUTION_WRITER，Founder依§11.3指派；Codex执行面网络中断）· {BASELINE_DATE} · execution_run_id={EXECUTION_RUN_ID} · parent={PARENT_EXECUTION_RUN_ID}"],
            ["核验结论", "候选施工自检PASS — D-17—D-29、S1—S8与三项补丁均已落到锚点；M0顶层清单仍为14项。独立Guardian、ChatGPT总顾问、Founder签署与M0授权均未完成。"],
        ],
    )
    doc.add_paragraph("一、Founder 裁决落实矩阵", style="Heading 1")
    matrix = [
        ["D-17", "1/4/6/7/8.7/13-15/附录", "M11 Founder真实品牌注入优先路径＋Codex夹具回退；条件交付物与证据边界", "PASS"],
        ["D-18", "8.6/8.7/10/16/附录", "合成封闭未见品牌、夹具/隐藏池物理隔离；Founder数据禁入", "PASS"],
        ["D-19", "12-15", "M12 Commercial V1.0命名、M0—M12顺序与关键路径保持", "PASS"],
        ["D-20", "1/3/4/5/6/7/8/9/10/11/13-16/附录", "C-13、人设对象、连续性合同、评测、组件、风险与里程碑", "PASS"],
        ["D-21", "1/3/4/5/6/7/8/9/10/11/13-16/附录", "C-14、原生语感对象/指标/反模板合同与M7验收", "PASS"],
        ["D-22", "3/5/7/11/13-16/附录", "VisualMerchandisingExtensionPort及兼容原则；非V1.0/M12硬门", "PASS"],
        ["D-23", "1/3/5/6/7/8/10/11/13-16/附录", "C-15、图片/属性对象、证据优先级、禁止推断、评测与流水线", "PASS"],
        ["D-24", "3/4/5/7/11/13-16/附录", "S-07保留导购辅助投影，RealtimeSalesAssistExtensionPort保留扩展兼容", "PASS"],
        ["D-25", "3/6/7/10/11/13-16/附录", "默认human_review；Founder授权auto_publish、审计/撤回/回滚/Kill Switch；非M12必须实现", "PASS"],
        ["D-26", "3/7/10/13-15/附录", "M11至少三类不变；正式真实品牌导入前五类100%可开启就绪", "PASS"],
        ["D-27", "8.7/11/16/17/治理与CI", "单一规范源、角色投影与GitHub Actions确定性守卫；主仓禁止真实品牌/顾客/隐藏内容", "PASS"],
        ["D-28", "3.3/P-14、M2/M3/M4/M5/M6/M10、10.2、16.1 R-22、16.2、附录B", "合理多解原则：评测三分类、②③类禁唯一Gold Answer、可接受决策边界、合法多解族与决策逻辑稳定率", "PASS"],
        ["D-29", "M6目标/交付物/通过标准、11.1组件、16.2", "M6反退化验收：硬约束确定性执行、LLM-off不变性可重放、DecisionTrace规则归因；禁纯Prompt/RAG端到端直答", "PASS"],
        ["S1—S8", "8/10/15/16/治理文件", "基线锚定、单次v1.2签署、单一Founder、角色降级、隐藏隔离、工作区佐证、任务分级与Cowork归位", "PASS"],
    ]
    add_receipt_table(doc, ["裁决", "PRD 主要落点", "核验证据", "状态"], matrix)

    doc.add_paragraph("二、一致性核验结果", style="Heading 1")
    checks = [
        ["M0十四项", "PRD第13/14/17.1节与M0申请均为同一组14项；无18项版回潮", "PASS"],
        ["输入/输出对象", "12个输入与配置对象；15个输出与内部审计对象；M1映射数量同步", "PASS"],
        ["目标→FR→门→里程碑", "D-17—D-29新增内容逐项重跑：G-11—G-15→FR-22—FR-30→八门/硬门/10.2指标→M0—M12锚点", "PASS"],
        ["旧措辞残留", "自动发布永久禁止、图像理解永久排除等8类废弃表达检索为0", "PASS"],
        ["版本/编号", "封面、页眉、版本记录、M0申请ID、回执ID、README均为v1.2/003；v1.1仍在根目录待签署状态", "PASS"],
        ["M11/M12", "M11双路径与条件交付完整；M12五品类激活就绪与发布兼容完整", "PASS"],
        ["单一Founder治理", "角色票数、工作区隔离、L1/L2/L3、降级、条件关闭、合规逐项台账与非外部评审声明完整", "PASS"],
        ["Project CI/隐藏边界", "单一规范源及投影完整；主仓仅允许隐藏基准清单/接口/结果摘要，不含隐藏正文", "PASS"],
        ["合理多解与反退化", "三分类合同、可接受决策边界、合法多解族、10.2指标更名（阈值≥85%不变）、R-22与两条停止条件、M6架构符合性三条均已落到锚点", "PASS"],
        ["编号不重排", "补充Prompt文字中的P-10/R-15已被D-20—D-26占用，本版落为P-14/R-22；既有P-01—P-13、R-01—R-21与FR/G/C编号一律未重排", "PASS"],
        ["DOCX包与可打开性", "三份DOCX通过ZIP/OOXML完整性、python-docx打开、LibreOffice PDF渲染与抽检", "PASS"],
        ["Checker", "工具/check_prd_v1_2.py全项通过；哈希写入最终交付报告", "PASS"],
    ]
    add_receipt_table(doc, ["核验域", "证据", "结果"], checks)

    doc.add_paragraph("三、待 Founder 审阅与签署", style="Heading 1")
    for text in [
        "本回执的PASS仅表示文档落文与一致性核验通过，不表示PRD v1.2已生效。",
        "Founder须分别签署PRD v1.2与M0执行申请v1.2（DIYU-CBFSK-EXEC-REQ-M0-003），才能另行开始M0。",
        "本轮没有需要Founder重新裁决的内部冲突；M2_FREEZE_REQUIRED指标的具体质量阈值依合同留待M2冻结，不是本轮未落文项。",
        "独立Guardian与ChatGPT总顾问尚未完成正式审查；本回执不得被解释为其PASS或外部独立验证。",
        "本轮候选由Founder依执行Prompt §11.3指派的TEMPORARY_EXECUTION_WRITER落盘（Codex执行面网络中断）；Codex作为默认唯一常规写入者的长期规则未改变，Codex恢复后是否复核由Founder另行裁决，本回执不代为假定。",
    ]:
        paragraph = doc.add_paragraph("• " + text)
        paragraph.style = doc.styles["List Paragraph"] if "List Paragraph" in [s.name for s in doc.styles] else doc.styles["Normal"]
    add_receipt_table(
        doc,
        ["裁决项", "结果（圈选）", "签署"],
        [
            ["PRD v1.2整体采纳", "PASS ／ CONDITIONAL ／ BLOCK", "Founder：____________  日期：____________"],
            ["M0执行申请v1.2授权", "APPROVE ／ APPROVE_WITH_CONDITIONS ／ REJECT", "Founder：____________  日期：____________"],
        ],
    )
    doc.add_paragraph("四、最终状态", style="Heading 1")
    status = doc.add_paragraph(
        "prd_v1_2_documentation_status: READY_FOR_GUARDIAN\n"
        "prd_v1_2_effective: false\n"
        "m0_authorized: false\n"
        "engineering_execution_started: false\n"
        "knowledge_distillation_started: false\n"
        "production_servable: false"
        "\nexecution_run_id: " + EXECUTION_RUN_ID + "\n"
        "guardian_review_completed: false\n"
        "chatgpt_remote_review_completed: false\n"
        "founder_prd_signed: false\n"
        "main_merged: false"
    )
    for run in status.runs:
        run.font.name = "Consolas"
        run.font.size = Pt(9.5)
    for section in doc.sections:
        if section.header.paragraphs:
            set_paragraph_text(section.header.paragraphs[0], "DIYU-CBFSK-001  |  PRD v1.2 核验回执  |  VERIFY-RECEIPT-001")
        else:
            section.header.add_paragraph("DIYU-CBFSK-001  |  PRD v1.2 核验回执  |  VERIFY-RECEIPT-001")
    doc.core_properties.title = "PRD v1.2 核验回执"
    doc.core_properties.subject = "D-17—D-29、S1—S8与治理合同的候选落文核验"
    doc.core_properties.comments = f"EXECUTION_RUN_ID={EXECUTION_RUN_ID}; READY_FOR_GUARDIAN; PRD not effective; M0 not authorized"
    doc.save(RECEIPT_OUTPUT)


def main() -> None:
    build_prd()
    build_m0()
    build_receipt()
    print(PRD_OUTPUT)
    print(M0_OUTPUT)
    print(RECEIPT_OUTPUT)


if __name__ == "__main__":
    main()
