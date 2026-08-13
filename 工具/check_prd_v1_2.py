#!/usr/bin/env python3
"""Contract checker for the PRD v1.2 documentation baseline."""

from __future__ import annotations

import argparse
import importlib.util
import re
import sys
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parent.parent
PRD = ROOT / "笛语跨品牌服装搭配专家内核_PRD与执行里程碑_v1.2.docx"
M0 = ROOT / "笛语跨品牌服装搭配专家内核_M0执行申请_v1.2.docx"
RECEIPT = ROOT / "PRD_v1.2_核验回执.docx"
README = ROOT / "README.md"
CHANGE_MAP = ROOT / "PRD_v1.2_change_map.yaml"


FORBIDDEN = [
    "V1.0全生命周期保持人工审核在环",
    "所有商业发布门通过后仍不得自动发布",
    "图像识别属性提取不属于V1.0",
    "商品属性只能来自PIM/ERP，图片不得参与",
    "完整陈列能力永久退出",
    "实时导购能力永久退出",
    "人设仅由账号语气合同承担",
    "叙事质量等同于自媒体语感",
]



class Checker:
    def __init__(self) -> None:
        self.errors: list[str] = []
        self.passes: list[str] = []

    def check(self, condition: bool, label: str, detail: str = "") -> None:
        if condition:
            self.passes.append(label)
        else:
            self.errors.append(label + (f": {detail}" if detail else ""))

    def contains_all(self, text: str, tokens: list[str], label: str) -> None:
        missing = [token for token in tokens if token not in text]
        self.check(not missing, label, f"missing={missing}")


def document_text(doc: Document) -> str:
    chunks = [paragraph.text for paragraph in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            chunks.extend(cell.text for cell in row.cells)
    for section in doc.sections:
        chunks.extend(paragraph.text for paragraph in section.header.paragraphs)
        chunks.extend(paragraph.text for paragraph in section.footer.paragraphs)
    return "\n".join(chunks)


def table_by_second_row(doc: Document, value: str):
    return next(table for table in doc.tables if len(table.rows) > 1 and table.cell(1, 0).text.strip() == value)


def ids_from_table(doc: Document, first_id: str, prefix: str) -> list[int]:
    table = table_by_second_row(doc, first_id)
    result: list[int] = []
    for row in table.rows[1:]:
        match = re.fullmatch(rf"{re.escape(prefix)}-(\d+)", row.cells[0].text.strip())
        if match:
            result.append(int(match.group(1)))
    return result


def heading_region(doc: Document, start: str, end: str) -> list[str]:
    paragraphs = doc.paragraphs
    start_index = next(i for i, paragraph in enumerate(paragraphs) if paragraph.text.strip() == start)
    end_index = next(i for i, paragraph in enumerate(paragraphs) if paragraph.text.strip() == end and i > start_index)
    return [paragraph.text.strip() for paragraph in paragraphs[start_index + 1 : end_index]]





def check_m0_lists(checker: Checker) -> None:
    """M0 十四项只有一个实现：ci/checkers/check_m0_fourteen_items.py。本文件不再自行提取。"""
    spec = importlib.util.spec_from_file_location(
        "check_m0_fourteen_items", ROOT / "ci" / "checkers" / "check_m0_fourteen_items.py"
    )
    module = importlib.util.module_from_spec(spec)
    sys.path.insert(0, str(ROOT / "ci" / "checkers"))
    try:
        spec.loader.exec_module(module)
        errors = module.validate(module.collect())
    finally:
        sys.path.remove(str(ROOT / "ci" / "checkers"))
    checker.check(not errors, "M0 fourteen items (delegated to ci/checkers)", "; ".join(errors))


def lifecycle_state() -> tuple[bool, bool]:
    """(v1.2 已生效, 活基线已切换到 v1.2) —— 单一真源：change map 的 resulting_state。

    此前用 --require-archive 这个 CLI 开关表达同一件事，等于把生命周期状态写成两处；
    现在只从状态派生，开关退化为「断言状态确实已到该阶段」。
    """
    import yaml

    state = yaml.safe_load(CHANGE_MAP.read_text(encoding="utf-8"))["resulting_state"]
    return bool(state["prd_v1_2_effective"]), state["current_active_baseline"] == "PRD_v1.2"


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument(
        "--require-archive",
        action="store_true",
        help="断言当前状态确实已完成 v1.1 归档；状态未到该阶段则判失败。",
    )
    args = parser.parse_args()
    checker = Checker()
    effective, switched = lifecycle_state()

    for path in (PRD, M0, RECEIPT, README, CHANGE_MAP):
        checker.check(path.exists(), f"file exists: {path.name}")
    if checker.errors:
        for error in checker.errors:
            print("FAIL", error)
        sys.exit(1)

    prd = Document(PRD)
    m0 = Document(M0)
    receipt = Document(RECEIPT)
    prd_text = document_text(prd)
    m0_text = document_text(m0)
    receipt_text = document_text(receipt)
    readme_text = README.read_text(encoding="utf-8")
    map_text = CHANGE_MAP.read_text(encoding="utf-8")

    checker.contains_all(
        prd_text,
        [
            "PRD v1.2 · 人设连续性、多模态商品理解与扩展兼容基线",
            "PROJECT_INITIATED / EXECUTION_NOT_STARTED / M0_AUTHORIZED_FALSE",
            "production_servable",
            "待Founder签署生效",
        ],
        "document status and version",
    )
    headers = "\n".join(paragraph.text for section in prd.sections for paragraph in section.header.paragraphs)
    checker.check("PRD v1.2" in headers and "PRD v1.0" not in headers and "PRD v1.1" not in headers, "PRD header version is current", headers)
    for phrase in FORBIDDEN:
        checker.check(phrase not in prd_text, f"forbidden phrase absent: {phrase}")

    checker.check(ids_from_table(prd, "G-01", "G") == list(range(1, 16)), "G IDs are consecutive G-01..G-15")
    checker.check(ids_from_table(prd, "C-01", "C") == list(range(1, 16)), "C IDs are consecutive C-01..C-15")
    checker.check(ids_from_table(prd, "FR-01", "FR") == list(range(1, 31)), "FR IDs are consecutive FR-01..FR-30")
    checker.check(ids_from_table(prd, "NFR-01", "NFR") == list(range(1, 13)), "NFR IDs are consecutive NFR-01..NFR-12")
    checker.check(ids_from_table(prd, "R-01", "R") == list(range(1, 23)), "risk IDs are consecutive R-01..R-22")
    checker.check(ids_from_table(prd, "P-01", "P") == list(range(1, 15)), "principle IDs are consecutive P-01..P-14")

    inputs = table_by_second_row(prd, "UniversalExpertKernel")
    outputs = table_by_second_row(prd, "BrandAudienceInterpretation")
    checker.check(len(inputs.rows) - 1 == 12, "input/configuration object count is 12", str(len(inputs.rows) - 1))
    checker.check(len(outputs.rows) - 1 == 15, "output/audit object count is 15", str(len(outputs.rows) - 1))
    checker.contains_all(prd_text, ["ProductImageAssetBundle", "StylistPersonaProfile", "PersonaMemorySnapshot", "FounderProvidedRealBrandPackage", "PublicationPolicy"], "new input/configuration objects")
    checker.contains_all(prd_text, ["VisualAttributeExtractionResult", "PersonaContinuityUpdate", "SocialMediaVoicePlan", "PublicationDecision"], "new output/audit objects")

    checker.contains_all(prd_text, ["C-13", "Stylist Persona Continuity Intelligence", "C-14", "Social-Media Native Voice Intelligence", "C-15", "Multimodal Garment Understanding"], "formal capabilities C-13..C-15")
    checker.contains_all(prd_text, ["VisualMerchandisingExtensionPort", "StoreSpaceContext", "PlanogramContext", "RealtimeSalesAssistExtensionPort", "SalesAssociateSessionContext", "RealtimeRecommendationResult"], "extension-compatible contracts")
    checker.contains_all(prd_text, ["V1.0 默认人工审核在环", "Founder 按租户/品牌/账号/风险级别显式授权", "审计、撤回与 Kill Switch", "PublicationPolicyController", "AutoPublishKillSwitch", "unauthorized_auto_publish_rate = 0", "FR-17"], "Founder-controlled publication contract (D-25 patch)")
    checker.contains_all(prd_text, ["authoritative_structured_fact", "human_verified_visual_attribute", "multimodal_inferred_visual_attribute", "authoritative_fact_override_rate = 0", "unverifiable_function_claim_rate = 0"], "multimodal evidence grading and hard gates")

    fr_table = table_by_second_row(prd, "FR-01")
    for number in range(22, 30):
        row = next(row for row in fr_table.rows if row.cells[0].text.strip() == f"FR-{number}")
        acceptance = row.cells[3].text
        checker.check(all(token in acceptance for token in ("验收：", "失败状态：", "里程碑：")), f"FR-{number} has acceptance/failure/milestone trace", acceptance)

    checker.contains_all(prd_text, ["stylist_persona_profile.schema.v0.1.json", "product_image_asset_bundle.schema.v0.1.json", "publication_policy.schema.v0.1.json", "extension_port_contracts", "12个输入与配置对象", "15个输出与内部审计对象"], "M1 schema and object mapping")
    checker.contains_all(prd_text, ["persona_continuity_scoring_rubric", "social_media_native_voice_scoring_rubric", "multimodal_attribute_benchmark", "multimodal_confidence_calibration_contract", "five_category_readiness_definition", "M2_FREEZE_REQUIRED"], "M2 evaluation additions")
    checker.contains_all(prd_text, ["persona_consistency_checker", "social_media_native_voice_checker", "anti_template_checker", "persona_and_voice_qualification_report"], "M7 deliverables")
    checker.contains_all(prd_text, ["FounderProvidedRealBrandPackage", "FounderInjectedRealBrandProductionTest", "founder_provided_real_brand_package_manifest", "founder_real_brand_test_production_report", "multimodal_product_understanding_report", "founder_capability_assessment_receipt", "founder_real_brand_package_provided=false", "fallback_source=codex_synthetic_fixture"], "M11 dual paths and conditional evidence")
    checker.contains_all(prd_text, ["five_category_activation_readiness=100%", "five_category_activation_readiness_report", "persona_state_versioning_and_rollback", "auto_publish_kill_switch_contract", "extension_port_registry"], "M12 readiness and governance")

    checker.contains_all(prd_text, ["R-15", "R-16", "R-17", "R-18", "R-19", "R-20", "R-21"], "risk additions R-15..R-21")
    checker.contains_all(prd_text, ["多模态推断覆盖权威", "未经Founder授权启用自动发布", "人设虚构履历", "Founder注入数据进入隐藏集", "扩展模块破坏现有输出Schema"], "new stop conditions")

    required_terms = ["Stylist Persona Continuity", "Social-Media Native Voice", "Persona Memory Snapshot", "Published Viewpoint Ledger", "Multimodal Garment Understanding", "Visual Attribute Provenance", "FounderProvidedRealBrandPackage", "FounderInjectedRealBrandProductionTest", "Publication Policy", "Visual Merchandising Extension Port", "Realtime Sales Assist Extension Port", "Five-category Activation Readiness"]
    term_table = next(table for table in prd.tables if table.cell(0, 0).text.strip() == "术语")
    term_text = "\n".join(cell.text for row in term_table.rows for cell in row.cells)
    checker.contains_all(term_text, required_terms, "all new terms are in glossary")

    check_m0_lists(checker)
    checker.contains_all(m0_text, ["DIYU-CBFSK-EXEC-REQ-M0-003", "PRD v1.2", "M0不得实际接收", "M0不得运行多模态", "M0不得建立搭配师人设记忆生产库", "M0不得启用自动发布", "人设与语感闭环", "多模态事实边界", "扩展兼容"], "M0 v1.2 control and Guardian additions")
    checker.contains_all(receipt_text, ["DIYU-CBFSK-PRD-V1.2-VERIFY-RECEIPT-001", "D-17", "D-26", "D-28", "D-29", "S1—S8", "READY_FOR_GUARDIAN", "m0_authorized: false", "production_servable: false", "guardian_review_completed: false", "chatgpt_remote_review_completed: false"], "verification receipt identifiers and final state")

    readme_state_tokens = (
        ["SIGNED", "prd_v1_2_effective: true", "归档_v1.1/", "DIYU-CBFSK-FOUNDER-SIGNOFF-001"]
        if switched
        else ["PENDING_FOUNDER_SIGNATURE", "prd_v1_2_effective: true", "DIYU-CBFSK-FOUNDER-SIGNOFF-001"]
        if effective
        else ["PENDING_FOUNDER_SIGNATURE", "READY_FOR_GUARDIAN", "prd_v1_2_effective: false", "目前没有 `归档_v1.1/`"]
    )
    checker.contains_all(readme_text, readme_state_tokens, "README lifecycle state block")
    checker.contains_all(readme_text, ["当前活基线", "PRD v1.2", "DIYU-CBFSK-EXEC-REQ-M0-003", "Founder 真实品牌注入", "Codex 夹具合成回退", "多模态商品理解", "人设连续性", "自媒体原生语感", "VisualMerchandisingExtensionPort", "RealtimeSalesAssistExtensionPort", "five_category_activation_readiness=100%", "合理多解原则", "LLM-off 不变性"], "README current-baseline index")
    checker.contains_all(
        prd_text,
        [
            "Non-Uniqueness by Design（合理多解原则）",
            "evaluation_task_class_contract.v0.1.yaml",
            "acceptable_decision_boundary_registry.v0.1.yaml",
            "open_decision_question_template_contract.v0.1.yaml",
            "disagreement_classification_and_solution_family_ledger.v0.1.yaml",
            "legal_decision_space_conformance_report",
            "核心决策逻辑稳定率",
            "唯一答案偏误",
            "被设置唯一Gold Answer：停止M2冻结",
        ],
        "D-28 non-uniqueness anchors",
    )
    checker.contains_all(
        prd_text,
        [
            "architecture_conformance_check_report",
            "llm_off_invariance_replay_fixtures",
            "架构符合性检查三条全部通过",
            "不是唯一argmax（D-28）",
            "每个排除项须标注触发它的确定性规则",
            "M6以纯Prompt/RAG端到端直答实现",
            "LLM-off Invariance（LLM-off不变性）",
        ],
        "D-29 anti-degeneration anchors",
    )
    checker.check(
        prd_text.count("reviewer_calibration_contract.v0.1.yaml") == 2,
        "reviewer_calibration_contract appears once per milestone list (13/14), not duplicated",
        f"count={prd_text.count('reviewer_calibration_contract.v0.1.yaml')}",
    )

    checker.contains_all(map_text, ["D-17:", "D-26:", "D-27:", "D-28:", "D-29:", "governance_ruling_map:", "m0_top_level_deliverable_count: 14", "input_and_configuration_objects: 12", "output_and_internal_audit_objects: 15", "functional_requirements: 30", "risks: 22",
                          f"documentation_status: {'FOUNDER_SIGNED' if effective else 'READY_FOR_GUARDIAN'}"],
                         "machine-readable change map")

    archive = ROOT / "归档_v1.1"
    v1_1_files = [
        "笛语跨品牌服装搭配专家内核_PRD与执行里程碑_v1.1.docx",
        "笛语跨品牌服装搭配专家内核_M0执行申请_v1.1.docx",
        "PRD_v1.1_核验回执.docx",
    ]
    if args.require_archive and not switched:
        checker.check(False, "--require-archive used while current_active_baseline is not PRD_v1.2")
    if switched:
        checker.check(effective, "v1.1 archived only after PRD v1.2 is effective")
        checker.check(archive.is_dir(), "v1.1 archive directory exists")
        for name in v1_1_files:
            checker.check((archive / name).is_file(), f"archived: {name}")
            checker.check(not (ROOT / name).exists(), f"v1.1 removed from root: {name}")
    else:
        # 红线：v1.2 生效前不得归档 v1.1；活基线切换前 v1.1 必须留在根目录。
        checker.check(not archive.exists(), "v1.1 not archived before the active-baseline switch")
        for name in v1_1_files:
            checker.check((ROOT / name).is_file(), f"active baseline stays at root: {name}")

    for label in checker.passes:
        print("PASS", label)
    if checker.errors:
        for error in checker.errors:
            print("FAIL", error)
        print(f"RESULT FAIL: {len(checker.errors)} error(s), {len(checker.passes)} pass(es)")
        sys.exit(1)
    print(f"RESULT PASS: {len(checker.passes)} checks")


if __name__ == "__main__":
    main()
